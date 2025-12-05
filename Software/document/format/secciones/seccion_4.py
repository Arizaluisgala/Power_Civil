from docx.shared import Pt, Cm
import os
import sys
project_root = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..', '..'))
sys.path.append(project_root)
sys.path.append(os.path.join(project_root, 'scripts'))
from scripts.table_spectrum import generar_spectrums_from_excel


def generar_espectro(nuevo,idioma,excel_spectro_path):
        """
        Genera la sección de espectros de diseño en el documento Word.
        Pasandole un excel que recorre y extrar los datos de Te y aac.
        """
        
        nuevo.add_paragraph()  # Espacio inicial
    
        # Párrafo introductorio según idioma
        if idioma in ["ingles", "en"]:
            parrafo_espectro = (
                "Below are the design spectra in each direction, defined according to the type of structural system and the previously indicated values of Ω0, Cd, and R."
            )
        else:
            parrafo_espectro = (
                "A continuación se muestran los espectros de diseño en cada dirección, definidos acorde al tipo de sistema estructural y los valores de Ω0, Cd y R indicados previamente."
            )
        p_intro = nuevo.add_paragraph(parrafo_espectro)
        p_intro.alignment = 0  # Izquierda
        run_intro = p_intro.runs[0] if p_intro.runs else p_intro.add_run()
        run_intro.font.name = "Arial"
        run_intro.font.size = Pt(12)
        run_intro.bold = False

        # Generar o pedir las gráficas de espectro
        # Puedes pasar el idioma para que las leyendas salgan correctas
        espectro_imgs = generar_spectrums_from_excel(lang="en" if idioma in ["ingles", "en"] else "es", excel_file=excel_spectro_path)

        # Insertar todas las imágenes y leyendas una debajo de la otra en la misma página
        if espectro_imgs:
            for idx, (img_path, eje) in enumerate(espectro_imgs, 7):
                if idioma in ["ingles", "en"]:
                    leyenda = f"Figure {idx}. Modified design spectrum in {('X' if eje=='X' else 'Y' if eje=='Y' else 'Z' if eje=='Z' else '')} direction"
                else:
                    leyenda = f"Figura {idx}. Espectro de diseño modificado en dirección {('X' if eje=='X' else 'Y' if eje=='Y' else 'Z' if eje=='Z' else '')}"

                # Imagen centrada
                p_img = nuevo.add_paragraph()
                p_img.alignment = 1
                run_img = p_img.add_run()
                run_img.add_picture(img_path, width=Cm(15), height=Cm(8))  # Ajusta el tamaño según sea necesario

                # Leyenda centrada justo debajo
                p_leyenda = nuevo.add_paragraph(leyenda)
                p_leyenda.alignment = 1
                run_leyenda = p_leyenda.runs[0] if p_leyenda.runs else p_leyenda.add_run()
                run_leyenda.font.name = "Arial"
                run_leyenda.font.size = Pt(12)

                # Espacio entre imágenes
                if idx != len(espectro_imgs) + 7 - 1:
                    nuevo.add_paragraph()
            # No agregar salto de página, así todas las imágenes quedan en la misma página
        nuevo.add_paragraph()  # Espacio final

        # Forzar salto de página al final de la sección
        nuevo.add_page_break()
        return nuevo    