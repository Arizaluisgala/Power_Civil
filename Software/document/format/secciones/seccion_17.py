from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH



def viga_image_v(nuevo,image_slots, idioma):
    if idioma== 'en':

        figura = "Figure 19. Beam with the greatest vertical displacement"
    else:
        figura = "Figura 19. Viga con mayor desplazamiento vertical"
        
        
    if not image_slots:
        print("⚠️ No hay image_slots disponibles")
        return nuevo    

    # Imagen centrada, tamaño fijo
    image_path = image_slots.get(8, None)
    p_img = nuevo.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_img = p_img.add_run()
    run_img.add_picture(image_path, width=Cm(5), height=Cm(16))

    # Título de la figura debajo, centrado, Arial 12
    nuevo.add_paragraph()  # Espacio extras
    p_title = nuevo.add_paragraph(figura)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.runs[0]
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(12)

    nuevo.add_page_break()