from docx.shared import Pt
from docx.oxml import parse_xml
import pandas as pd
import os
import sys
project_root = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..', '..'))
sys.path.append(project_root)
sys.path.append(os.path.join(project_root, 'scripts'))
sys.path.append(os.path.join(project_root, 'componentes'))

from scripts.extract_tables_of_excel import extract_tables_from_excel_reacciones
from componentes.poner_bordes_tablas import poner_bordes_tabla
from componentes.set_repeat_header import set_repeat_table_header
from document.format.helpers import format_number


def reacciones_de_la_estructura(nuevo,idioma,file_path):
        nuevo.add_paragraph()
        if idioma in ["ingles", "en"]:
            nuevo.add_paragraph("Reactions at the base of the structure")
            p = nuevo.paragraphs[-1]
            run = p.runs[0] if p.runs else p.add_run()
            run.font.name = "Arial"
            run.font.size = Pt(12)
            run.bold = True
            p.alignment = 0
            nuevo.add_paragraph()
            parrafo = "For the design of the structure's foundations, the reactions for the service and ultimate load cases are obtained (200 and 1200)."
        else:
            nuevo.add_paragraph("Reacciones en la base de la estructura")
            p = nuevo.paragraphs[-1]
            run = p.runs[0] if p.runs else p.add_run()
            run.font.name = "Arial"
            run.font.size = Pt(12)
            run.bold = True
            p.alignment = 0
            nuevo.add_paragraph()
            parrafo = "Para el diseño de las cimentaciones de la estructura se obtienen las reacciones para los casos de carga de servicio y últimas (200 y 1200)."
        p_intro = nuevo.add_paragraph(parrafo)
        p_intro.alignment = 0
        run_intro = p_intro.runs[0] if p_intro.runs else p_intro.add_run()
        run_intro.font.name = "Arial"
        run_intro.font.size = Pt(12)
        nuevo.add_paragraph()

        # Extraer tabla de reacciones
        try:
            reacciones_dict = extract_tables_from_excel_reacciones(file_path)
            df_reacciones = reacciones_dict.get('R', None)
        except Exception as e:
            df_reacciones = None

        HEADERS_ES = [
            "Nodo",
            "Caso de carga",
            "FX (kN)",
            "FY (kN)",
            "FZ (kN)",
            "MX \n(kN-m)",
            "MY \n(kN-m)",
            "MZ \n(kN-m)",
            "Componente Crítico"
        ]
        HEADERS_EN = [
            "Node",
            "Load Case",
            "FX (kN)",
            "FY (kN)",
            "FZ (kN)",
            "MX \n(kN-m)",
            "MY \n(kN-m)",
            "MZ \n(kN-m)",
            "Critical Component"
        ]
        COL_MAP = {
            "Nodo": ["Nodo", "Node"],
            "Node": ["Node", "Nodo"],
            "Caso de carga": ["Caso de carga", "Caso", "Load Case"],
            "Load Case": ["Load Case", "Caso de carga", "Caso"],
            "FX (kN)": ["FX (kN)"],
            "FY (kN)": ["FY (kN)"],
            "FZ (kN)": ["FZ (kN)"],
            "MX \n(kN-m)": ["MX (kN-m)", "MX"],
            "MY \n(kN-m)": ["MY (kN-m)", "MY"],
            "MZ \n(kN-m)": ["MZ (kN-m)", "MZ"],
            "Componente Crítico": ["Componente Crítico", "Critical Component"],
            "Critical Component": ["Critical Component", "Componente Crítico"]
        }

        if df_reacciones is not None and not df_reacciones.empty:
            headers = HEADERS_EN if idioma in ["ingles", "en"] else HEADERS_ES
            from docx.shared import Cm
            from docx.enum.table import WD_TABLE_ALIGNMENT
            # Ajuste de ancho personalizado para columna Componente Crítico/Critical Component
            col_widths = []
            for h in headers:
                if h == "Componente Crítico" or h == "Critical Component":
                    col_widths.append(2.99)
                else:
                    col_widths.append((19.1 - 2.99) / (len(headers)-1))
            # Mapeo robusto de columnas
            cols = []
            for h in headers:
                found = None
                for c in COL_MAP[h]:
                    if c in df_reacciones.columns:
                        found = c
                        break
                cols.append(found)
            valid_cols = [c for c in cols if c is not None]
            df_show = df_reacciones[valid_cols].copy()
            df_show.columns = headers[:len(valid_cols)]
            for h in headers:
                if h not in df_show.columns:
                    df_show[h] = ""
            df_show = df_show[headers]
            df_show = df_show.dropna(how='all').copy()
            df_show = df_show[~df_show.apply(lambda row: all((str(x).strip()=="" or pd.isna(x)) for x in row), axis=1)]
            n = len(headers)
            tabla = nuevo.add_table(rows=1, cols=n)
            set_repeat_table_header(tabla.rows[0])
            tabla.autofit = False
            tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
            tabla.style = None
            # Encabezados
            for j, h in enumerate(headers):
                cell = tabla.cell(0, j)
                # Saltos de línea en encabezados
                cell.text = h.replace("\\n", "\n")
                p0 = cell.paragraphs[0]; p0.alignment = 1
                while len(p0.runs) > 1:
                    p0.runs[-1].clear(); del p0.runs[-1]
                run = p0.runs[0]
                run.text = cell.text
                run.font.name = "Arial"
                run.font.size = Pt(10)
                run.font.bold = True
                rPr = run._element.get_or_add_rPr()
                sz = parse_xml(r'<w:sz xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="20"/>')
                szCs = parse_xml(r'<w:szCs xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="20"/>')
                rPr.append(sz); rPr.append(szCs)
                shade = parse_xml(r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="D9D9D9"/>')
                cell._tc.get_or_add_tcPr().append(shade)
                if j < len(col_widths):
                    cell.width = Cm(col_widths[j])
                cell.vertical_alignment = 1
            tabla.rows[0].height = Cm(1.6)
            tabla.rows[0].height_rule = 2
            # Filas de datos
            for row in df_show.itertuples(index=False):
                cells = tabla.add_row().cells
                for j, val in enumerate(row):
                    text = ""
                    if pd.notna(val):
                        if isinstance(val, (int, float)):
                            text = format_number(val)
                        else:
                            text = str(val).strip()
                    cell = cells[j]
                    cell.text = text
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = 1
                        for run in paragraph.runs:
                            run.font.name = "Arial"
                            run.font.size = Pt(11)
                    if j < len(col_widths):
                        cell.width = Cm(col_widths[j])
                    cell.vertical_alignment = 1
                tabla.rows[-1].height = Cm(0.58)
                tabla.rows[-1].height_rule = 2
            poner_bordes_tabla(tabla)
            nuevo.add_paragraph()
