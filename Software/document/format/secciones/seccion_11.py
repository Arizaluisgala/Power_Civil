from docx.shared import Pt
from docx.oxml import parse_xml
import pandas as pd
import os
import sys
project_root = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..', '..'))
sys.path.append(project_root)
sys.path.append(os.path.join(project_root, 'scripts'))
sys.path.append(os.path.join(project_root, 'componentes'))

from scripts.extract_tables_of_excel import extract_tables_from_excel_ratios
from componentes.poner_bordes_tablas import poner_bordes_tabla
from componentes.set_repeat_header import set_repeat_table_header
from document.format.helpers import format_number


def ratios_tables(nuevo,idioma,file_path):  
        nuevo.add_paragraph()
        # Añado título en español o inglés según el idioma
        # Título y párrafo introductorio según idioma
        if idioma in ["ingles", "en"]:
            nuevo.add_paragraph("Verification of limit states of exhaustion by strength")
            p = nuevo.paragraphs[-1]
            run = p.runs[0] if p.runs else p.add_run()
            run.font.name = "Arial"
            run.font.size = Pt(12)
            run.bold = True
            p.alignment = 0
            nuevo.add_paragraph()  # Salto de línea
            parrafo = "The stress ratio in all members of the structure must be less than 1; the ratios for each member are shown in the following figure and/or table:"
        else:
            nuevo.add_paragraph("Verificación de estados límites de agotamiento por resistencia")
            p = nuevo.paragraphs[-1]
            run = p.runs[0] if p.runs else p.add_run()
            run.font.name = "Arial"
            run.font.size = Pt(12)
            run.bold = True
            p.alignment = 0
            nuevo.add_paragraph()  # Salto de línea
            parrafo = "El ratio de esfuerzos en todos los miembros de la estructura debe ser inferior a 1, en la siguiente figura y/o tabla se muestran los ratios para cada miembro:"

        # Agregar el párrafo introductorio en ambos idiomas
        p_intro = nuevo.add_paragraph(parrafo)
        p_intro.alignment = 0
        run_intro = p_intro.runs[0] if p_intro.runs else p_intro.add_run()
        run_intro.font.name = "Arial"
        run_intro.font.size = Pt(12)
            
        nuevo.add_paragraph()    
        
        # Tabla de ratios de esfuerzos
        ratios_dict = extract_tables_from_excel_ratios(file_path)
        df_ratios = ratios_dict.get('Ratios', None)

        # Encabezados en ambos idiomas
        HEADERS_ES = [
            "ID del Elemento", "Perfil", "Código de diseño", "Caso de carga crítico",
            "Sección crítica (m)", "Ratio crítico", "Estado de diseño"
        ]
        HEADERS_EN = [
            "Element ID", "Profile", "Design Code", "Critical Load Case",
            "Critical Section (m)", "Critical Ratio", "Design State"
        ]
        # Mapeo de columnas del excel a los encabezados (acepta variantes de mayúsculas/minúsculas)
        COL_MAP = {
            "ID del Elemento": ["ID del Elemento", "ID del elemento", "Element ID", "ID Elemento"],
            "Element ID": ["Element ID", "ID del Elemento", "ID del elemento", "ID Elemento"],
            "Perfil": ["Perfil", "Profile"],
            "Profile": ["Profile", "Perfil"],
            "Código de diseño": ["Código de diseño", "Código de Diseño", "Design Code"],
            "Design Code": ["Design Code", "Código de diseño", "Código de Diseño"],
            "Caso de carga crítico": ["Caso de carga crítico", "Caso de carga Crítico", "Critical Load Case"],
            "Critical Load Case": ["Critical Load Case", "Caso de carga crítico", "Caso de carga Crítico"],
            "Sección crítica (m)": ["Sección crítica (m)", "Sección Crítica (m)", "Critical Section (m)"],
            "Critical Section (m)": ["Critical Section (m)", "Sección crítica (m)", "Sección Crítica (m)"],
            "Ratio crítico": ["Ratio crítico", "Critical Ratio", "Ratio"],
            "Critical Ratio": ["Critical Ratio", "Ratio crítico", "Ratio"],
            "Estado de diseño": ["Estado de diseño", "Estado de Diseño", "Design State"],
            "Design State": ["Design State", "Estado de diseño", "Estado de Diseño"]
        }

        if df_ratios is not None and not df_ratios.empty:
            headers = HEADERS_EN if idioma in ["ingles", "en"] else HEADERS_ES
            col_widths = [2.42, 2.29, 2.29, 2.52, 2.05, 1.66, 3]
            # Mapeo robusto de columnas
            cols = []
            for h in headers:
                found = None
                for c in COL_MAP[h]:
                    if c in df_ratios.columns:
                        found = c
                        break
                cols.append(found)
            valid_cols = [c for c in cols if c is not None]
            df_show = df_ratios[valid_cols].copy()
            df_show.columns = headers[:len(valid_cols)]
            for h in headers:
                if h not in df_show.columns:
                    df_show[h] = ""
            df_show = df_show[headers]
            df_show = df_show.dropna(how='all').copy()
            df_show = df_show[~df_show.apply(lambda row: all((str(x).strip()=="" or pd.isna(x)) for x in row), axis=1)]

            n = len(headers)
            tabla = nuevo.add_table(rows=1, cols=n)
            set_repeat_table_header(tabla.rows[0])
            tabla.autofit = False
            tabla.style = None
            from docx.enum.table import WD_TABLE_ALIGNMENT
            tabla.alignment = WD_TABLE_ALIGNMENT.CENTER

            for j, h in enumerate(headers):
                cell = tabla.cell(0, j)
                cell.text = h
                p0 = cell.paragraphs[0]; p0.alignment = 1
                while len(p0.runs) > 1:
                    p0.runs[-1].clear(); del p0.runs[-1]
                run = p0.runs[0]
                run.text = h; run.font.name = "Arial"; run.font.size = Pt(11); run.font.bold = True
                rPr = run._element.get_or_add_rPr()
                sz = parse_xml(r'<w:sz xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="22"/>')
                szCs = parse_xml(r'<w:szCs xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="22"/>')
                rPr.append(sz); rPr.append(szCs)
                shade = parse_xml(r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="D9D9D9"/>')
                cell._tc.get_or_add_tcPr().append(shade)
                cell.width = Pt(col_widths[j]*28.35); cell.vertical_alignment = 1
            tabla.rows[0].height = Pt(1.6*28.35)

            for row in df_show.itertuples(index=False):
                cells = tabla.add_row().cells
                for j, val in enumerate(row):
                    cell = cells[j]
                    text = ""
                    if pd.notna(val):
                        if isinstance(val, (int, float)):
                            text = format_number(val)
                        else:
                            text = str(val).strip()
                    cell.text = text
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = 1
                        for run in paragraph.runs:
                            run.font.name = "Arial"
                            run.font.size = Pt(11)
                    # Estado de diseño: color (PASS en verde, FAIL en rojo)
                    if headers[j] in ["Estado de diseño", "Design State"]:
                        estado = text.lower().strip()
                        if estado in ["pass", "sí", "si", "yes", "aprobado", "aprobar", "aprobada", "aprobada"]:
                            cell.text = "PASS" if idioma in ["ingles", "en"] else "APROBADO"
                            for paragraph in cell.paragraphs:
                                paragraph.alignment = 1
                                for run in paragraph.runs:
                                    run.text = cell.text
                                    run.font.name = "Arial"
                                    run.font.size = Pt(11)
                                    run.bold = False
                            shade = parse_xml(r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="C6EFCE"/>')
                            cell._tc.get_or_add_tcPr().append(shade)
                        else:
                            cell.text = "FAIL" if idioma in ["ingles", "en"] else "FALLA"
                            for paragraph in cell.paragraphs:
                                paragraph.alignment = 1
                                for run in paragraph.runs:
                                    run.text = cell.text
                                    run.font.name = "Arial"
                                    run.font.size = Pt(10)
                                    run.bold = False
                            shade = parse_xml(r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="FFC7CE"/>')
                            cell._tc.get_or_add_tcPr().append(shade)
                    cell.width = Pt(col_widths[j]*28.35); cell.vertical_alignment = 1
                tabla.rows[-1].height = Pt(0.58*28.35)

            poner_bordes_tabla(tabla)
            nuevo.add_paragraph()
        else:
            if idioma in ["ingles", "en"]:
                nuevo.add_paragraph("No ratio verification data found.")
            else:
                nuevo.add_paragraph("No se encontraron datos de verificación de ratios.")
            nuevo.add_paragraph()
                
        # Salto de línea
        nuevo.add_paragraph()        
        
