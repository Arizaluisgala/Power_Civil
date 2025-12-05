from docx.shared import Pt, Cm
from docx.oxml import parse_xml
from docx.enum.table import WD_ROW_HEIGHT_RULE
import os
import sys
project_root = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..', '..'))
sys.path.append(project_root)
sys.path.append(os.path.join(project_root, 'scripts'))
sys.path.append(os.path.join(project_root, 'componentes'))

from scripts.extract_tables_of_excel import extract_tables_from_excel_miembros_fisicos
from scripts.extract_tables_of_excel import extract_tables_from_excel_materiales
from scripts.extract_tables_of_excel import extract_tables_from_excel_soportes_asignados_a_la_estructura
from componentes.poner_bordes_tablas import poner_bordes_tabla
from componentes.set_repeat_header import set_repeat_table_header
from document.format.helpers import format_number

# Diccionario de traducción para materiales
traduccion_materiales = {
    "Módulo de Young \"E\" (kN/m²)": "Young's Modulus \"E\" (kN/m²)",
    "Coeficiente de Poisson \"v\"": "Poisson's Ratio \"v\"",
    "Densidad \"Y\" (kN/m³)": "Density \"Y\" (kN/m³)",
    "Coeficiente Térmico \"a\"(/°C)": "Thermal Coefficient \"a\"(/°C)",
    "Amortiguamiento Crítico": "Critical Damping",
    "Módulo de Corte \"G\" (kN/m²)": "Shear Modulus \"G\" (kN/m²)",
    "Límite Elástico \"Fy\" (kN/m²)": "Yield Limit \"Fy\" (kN/m²)",
    "Resistencia a la Tracción \"Fu\" (kN/m²)": "Tensile Strength \"Fu\" (kN/m²)",
    "Relación de Resistencia a la Cedencia \"Ry\"": "Yield Strength Ratio \"Ry\"",
    "Relación de Resistencia a la Tracción \"Rt\"": "Tensile Strength Ratio \"Rt\"",
    "Resistencia a la Compresión Fcu (kN/m²)": "Compressive Strength Fcu (kN/m²)"
}

# Diccionario de traducción para soportes
traduccion_soportes = {
    "Soporte fijo con liberaciones": "Fixed support with releases",
    "Soporte impuesto con liberaciones": "Imposed support with releases"
}

def tablas_miembros_materiales_soportes(nuevo, idioma, file_path):
    """
    Genera la sección 2 del documento Word con:
    - Tabla de miembros físicos continuos y elementos analíticos
    - Tabla de propiedades de materiales
    - Tabla de soportes asignados a la estructura
    """

    print(f"🔍 DEBUG: Iniciando sección 2 con archivo: {file_path}")
    print(f"🔍 DEBUG: Idioma: {idioma}")
    print(f"🔍 DEBUG: Archivo existe: {os.path.exists(file_path)}")

    if not os.path.exists(file_path):
        print(f"❌ ERROR: No se encuentra el archivo Excel: {file_path}")
        nuevo.add_paragraph(f"Error: No se encontró el archivo Excel: {file_path}")
        return

    nuevo.add_paragraph()

    # --- Tabla de miembros físicos continuos y elementos analíticos ---
    print("🔍 DEBUG: Creando tabla de miembros físicos...")

    # Título principal
    if idioma in ["ingles", "en"]:
        nuevo.add_paragraph("Physical Member Table")
    else:
        nuevo.add_paragraph("Tabla de miembros físicos")

    p = nuevo.paragraphs[-1]
    run = p.runs[0] if p.runs else p.add_run()
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.bold = True
    p.alignment = 0

    nuevo.add_paragraph()

    try:
        miembros_fisicos_dict = extract_tables_from_excel_miembros_fisicos(file_path)
        pm_df = miembros_fisicos_dict.get('PM', None)
        if pm_df is not None and not pm_df.empty:
            if idioma in ["ingles", "en"]:
                titulo_tabla = "Numbering"
                headers = ["Physical Member (PM)", "Analytical Element (AE)"]
            else:
                titulo_tabla = "Numeración"
                headers = ["Miembros Físicos (MF)", "Elementos Analíticos (EA)"]

            ncols = 2
            nrows = len(pm_df) + 2  # +2: fila de título y encabezado

            table = nuevo.add_table(rows=nrows, cols=ncols)
            table.autofit = False
            table.alignment = 1  # Centrar la tabla
            # Asignar anchos fijos a columnas y celdas
            ancho_cols = [Cm(4), Cm(5.5)]
            for i, ancho in enumerate(ancho_cols):
                table.columns[i].width = ancho
                for row in table.rows:
                    row.cells[i].width = ancho

            # Fila de título fusionada y gris
            cell_title = table.cell(0, 0)
            cell_title.merge(table.cell(0, 1))
            cell_title.text = titulo_tabla
            p = cell_title.paragraphs[0]
            p.alignment = 1
            run = p.runs[0] if p.runs else p.add_run()
            run.font.bold = True
            run.font.size = Pt(12)
            run.font.name = "Arial"
            shading_elm = parse_xml(
                r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="D9D9D9"/>'
            )
            cell_title._tc.get_or_add_tcPr().append(shading_elm)
            table.rows[0].height = Cm(0.55)
            table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

            # Fila de encabezados (gris)
            for i, h in enumerate(headers):
                cell = table.cell(1, i)
                cell.text = h
                p = cell.paragraphs[0]
                p.alignment = 1
                run = p.runs[0] if p.runs else p.add_run()
                run.font.bold = True
                run.font.size = Pt(11)
                run.font.name = "Arial"
                shading_elm = parse_xml(
                    r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="D9D9D9"/>'
                )
                cell._tc.get_or_add_tcPr().append(shading_elm)
            set_repeat_table_header(table.rows[0])
            set_repeat_table_header(table.rows[1])
            table.rows[1].height = Cm(0.55)
            table.rows[1].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

            # Agregar filas de datos
            for idx, miembro in pm_df.iterrows():
                row = table.rows[idx + 2]
                val1 = miembro.get("ID Physical Member", miembro.get("Physical Member", ""))
                val2 = miembro.get("Elementos Analíticos (IDs)", miembro.get("Analytical Elements", ""))

                row.cells[0].text = str(val1)
                row.cells[1].text = str(val2)

                for i in range(2):
                    p = row.cells[i].paragraphs[0]
                    p.alignment = 1
                    run = p.runs[0] if p.runs else p.add_run()
                    run.font.size = Pt(10)
                    run.font.name = "Arial"
                    row.cells[i].width = ancho_cols[i]  # Forzar ancho de celda
                row.height = Cm(0.55)
                row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            poner_bordes_tabla(table)
        else:
            nuevo.add_paragraph("No data found" if idioma in ["ingles", "en"] else "No se encontraron datos")
    except Exception as e:
        print(f"❌ ERROR al extraer miembros físicos: {e}")
        nuevo.add_paragraph(f"Error: {str(e)}")

    # --- Tabla de soportes asignados a la estructura ---
    print("🔍 DEBUG: Creando tabla de soportes...")

    nuevo.add_paragraph()
    if idioma in ["ingles", "en"]:
        nuevo.add_paragraph("Table of supports assigned to the structure")
        headers = ["Reference", "Support", "Node"]
    else:
        nuevo.add_paragraph("Tabla de soportes asignados a la estructura")
        headers = ["Referencia", "Apoyo", "Nodo"]

    p = nuevo.paragraphs[-1]
    run = p.runs[0] if p.runs else p.add_run()
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.bold = True
    p.alignment = 0

    nuevo.add_paragraph()

    try:
        soportes_dict = extract_tables_from_excel_soportes_asignados_a_la_estructura(file_path)
        soportes_df = soportes_dict.get('Soporte', None)
        if soportes_df is not None and not soportes_df.empty:
            ncols = 3
            tabla_soportes = nuevo.add_table(rows=1, cols=ncols)
            set_repeat_table_header(tabla_soportes.rows[0])
            tabla_soportes.autofit = False
            tabla_soportes.alignment = 1  # Centrar la tabla
            # Asignar anchos fijos a columnas y celdas
            ancho_cols = [Cm(2.82), Cm(7.41), Cm(1.98)]
            for i, ancho in enumerate(ancho_cols):
                tabla_soportes.columns[i].width = ancho
                for row in tabla_soportes.rows:
                    row.cells[i].width = ancho

            # Fila de encabezados (gris)
            headers = ["Referencia", "Apoyo", "Nodo"] if idioma not in ["ingles", "en"] else ["Reference", "Support", "Node"]
            for i, h in enumerate(headers):
                cell = tabla_soportes.cell(0, i)
                cell.text = h
                p = cell.paragraphs[0]
                p.alignment = 1
                run = p.runs[0] if p.runs else p.add_run()
                run.font.bold = True
                run.font.size = Pt(11)
                run.font.name = "Arial"
                shading_elm = parse_xml(
                    r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="D9D9D9"/>'
                )
                cell._tc.get_or_add_tcPr().append(shading_elm)
                cell.width = ancho_cols[i]  # Forzar ancho de celda
            tabla_soportes.rows[0].height = Cm(0.55)
            tabla_soportes.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

            # Agregar filas de datos
            for _, row_data in soportes_df.iterrows():
                row = tabla_soportes.add_row()
                ref = row_data.get("Referencia", row_data.get("Reference", ""))
                tipo = row_data.get("Apoyo", row_data.get("Tipo de Soporte", row_data.get("Support Type", "")))
                if idioma in ["ingles", "en"]:
                    tipo = traduccion_soportes.get(str(tipo).strip(), str(tipo))
                nodo = row_data.get("Nodo", row_data.get("Node", ""))
                row.cells[0].text = str(ref)
                row.cells[1].text = str(tipo)
                row.cells[2].text = str(nodo)
                for i in range(3):
                    p = row.cells[i].paragraphs[0]
                    p.alignment = 1
                    run = p.runs[0] if p.runs else p.add_run()
                    run.font.size = Pt(10)
                    run.font.name = "Arial"
                    row.cells[i].width = ancho_cols[i]  # Forzar ancho de celda
                row.height = Cm(0.55)
                row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            poner_bordes_tabla(tabla_soportes)
        else:
            nuevo.add_paragraph("No support data found" if idioma in ["ingles", "en"] else "No se encontraron datos de soporte")
    except Exception as e:
        print(f"❌ ERROR al extraer soportes: {e}")
        nuevo.add_paragraph(f"Error loading support data: {str(e)}")

    # --- Título de la tabla de materiales ---
    if idioma in ["ingles", "en"]:
        nuevo.add_paragraph("Table of material properties")
    else:
        nuevo.add_paragraph("Tabla de propiedades de materiales")
    p = nuevo.paragraphs[-1]
    run = p.runs[0] if p.runs else p.add_run()
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.bold = True
    p.alignment = 0

    nuevo.add_paragraph()

    try:
        materiales_dict = extract_tables_from_excel_materiales(file_path)
        mc_df = materiales_dict.get('MC', None)
        if mc_df is not None and not mc_df.empty:
            # Detectar dinámicamente las columnas STEEL con datos
            steel_cols = [col for col in mc_df.columns if col.upper().startswith("STEEL")]
            # Filtrar solo las columnas STEEL que tengan al menos un dato no vacío
            steel_cols = [col for col in steel_cols if mc_df[col].notna().any() and (mc_df[col] != '').any()]
            # Encabezados según idioma
            if idioma in ["ingles", "en"]:
                headers = ["Property"] + steel_cols
                col_keys = ["Property"] + steel_cols
            else:
                headers = ["Propiedad"] + steel_cols
                col_keys = ["Título"] + steel_cols

            tabla_mat = nuevo.add_table(rows=1, cols=len(headers))
            set_repeat_table_header(tabla_mat.rows[0])
            tabla_mat.autofit = False
            tabla_mat.alignment = 1  # Centrar la tabla

            # Encabezados
            for i, h in enumerate(headers):
                cell = tabla_mat.cell(0, i)
                cell.text = h
                p = cell.paragraphs[0]
                p.alignment = 1
                p.runs[0].font.bold = True
                p.runs[0].font.size = Pt(11)
                shading_elm = parse_xml(
                    r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="D9D9D9"/>'
                )
                cell._tc.get_or_add_tcPr().append(shading_elm)
            tabla_mat.rows[0].height = Cm(0.55)
            tabla_mat.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

            # Datos
            for _, row_data in mc_df.iterrows():
                # Solo muestra filas que tengan algún dato en las columnas STEEL
                if any(str(row_data.get(col, '')).strip() for col in steel_cols):
                    table_row = tabla_mat.add_row()
                    # Primera columna: propiedad
                    if idioma in ["ingles", "en"]:
                        prop = row_data.get("Property", row_data.get("Título", ""))
                        prop = traduccion_materiales.get(prop, prop)
                    else:
                        prop = row_data.get("Título", row_data.get("Property", ""))
                    table_row.cells[0].text = str(prop)
                    # Columnas STEEL
                    for i, col in enumerate(steel_cols, start=1):
                        value = row_data.get(col, '')
                        try:
                            # Intentar convertir a float y formatear
                            value = float(value)
                            table_row.cells[i].text = format_number(value)
                        except (ValueError, TypeError):
                            # Si no se puede convertir, usar el valor original
                            table_row.cells[i].text = str(value)

                    # Formato y altura
                    for i in range(len(headers)):
                        p = table_row.cells[i].paragraphs[0]
                        p.alignment = 1
                        p.runs[0].font.size = Pt(10)
                    table_row.height = Cm(0.55)
                    table_row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

            poner_bordes_tabla(tabla_mat)
        else:
            nuevo.add_paragraph("No material data found" if idioma in ["ingles", "en"] else "No se encontraron datos de materiales")
    except Exception as e:
        print(f"❌ ERROR al extraer materiales: {e}")
        nuevo.add_paragraph(f"Error loading material data: {str(e)}")
