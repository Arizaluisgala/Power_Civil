import sys
import os
from docx.shared import Pt, Cm
from docx.oxml import parse_xml
from docx.enum.table import WD_ROW_HEIGHT_RULE
project_root = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..', '..'))
sys.path.append(project_root)
sys.path.append(os.path.join(project_root, 'staad_automation'))
sys.path.append(os.path.join(project_root, 'componentes'))
from componentes.poner_bordes_tablas import poner_bordes_tabla
from componentes.set_repeat_header import set_repeat_table_header
from document.format.helpers import format_number


def crear_tabla_word(doc, headers, data, col_widths):
    ncols = len(headers)
    nrows = len(data)
    tabla = doc.add_table(rows=nrows+1, cols=ncols)
    set_repeat_table_header(tabla.rows[0])
    tabla.autofit = False

    # Encabezados
    for j, h in enumerate(headers):
        cell = tabla.cell(0, j)
        cell.text = h
        p = cell.paragraphs[0]
        p.alignment = 1
        run = p.runs[0] if p.runs else p.add_run()
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.name = "Arial"
        shading_elm = parse_xml(
            r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="D9D9D9"/>'
        )
        cell._tc.get_or_add_tcPr().append(shading_elm)
        cell.width = Cm(col_widths[j])

    tabla.rows[0].height = Cm(0.8)
    tabla.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    # Filas de datos
    for i, row in enumerate(data):
        for j, val in enumerate(row):
            cell = tabla.cell(i+1, j)
            cell.text = str(val) if val is not None else ""
            p = cell.paragraphs[0]
            p.alignment = 1
            run = p.runs[0] if p.runs else p.add_run()
            run.font.size = Pt(10)
            run.font.name = "Arial"
            cell.width = Cm(col_widths[j])
        tabla.rows[i+1].height = Cm(0.8)
        tabla.rows[i+1].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    return tabla


def casos_de_carga_combinados(nuevo, idioma):
    nuevo.add_paragraph()
    parrafo_carga = nuevo.add_paragraph()
    if idioma in ["ingles", "en"]:
        parrafo_carga.add_run("According to the design criteria, the load combinations to be used in the analysis of the structure are the following:").font.name = "Arial"
    else:
        parrafo_carga.add_run("De acuerdo con los criterios de diseño, las combinaciones de cargas a utilizar en el análisis de la estructura son las siguientes:").font.name = "Arial"       

    nuevo.add_paragraph()
    
    try:
        from staad_automation.extract_load import extract_combinations_load
        # Obtener el excel desde variables locales del stack (como lo hace extract_combinations_load)
        import inspect
        caller_locals = inspect.currentframe().f_back.f_locals
        file_path = caller_locals.get('excel_file_path') or caller_locals.get('excel_file_path_cargas')
        combinaciones = extract_combinations_load(file_path)
    except Exception as e:
        print(f"⚠️ No se pudo extraer combinaciones de carga: {e}")
        combinaciones = None

    if combinaciones:
        titulos = {
            "Cargas de servicio": {
                "es": "Tabla de combinaciones de cargas para el estado límite de servicio",
                "en": "Load combination table for serviceability limit state"
            },
            "Cargas Ultimas": {
                "es": "Tabla de combinaciones de cargas para el estado de agotamiento resistente",
                "en": "Load combination table for ultimate limit state"
            },
            "Verificación con viento": {
                "es": "Tabla de combinaciones de cargas para verificación de desplazabilidad por viento",
                "en": "Load combination table for wind displacement verification"
            },
            "Verificación con sismo": {
                "es": "Tabla de combinaciones de cargas para verificación de desplazabilidad por sismo",
                "en": "Load combination table for seismic displacement verification"
            },
            "Conexiones": {
                "es": "Tabla de combinaciones de cargas para diseño de conexiones",
                "en": "Load combination table for connection design"
            }
        }

        from collections import defaultdict
        comb_por_tipo = defaultdict(list)
        for num, nombre, header in combinaciones:
            comb_por_tipo[header].append((num, nombre))

        # Ancho fijo: columna 1 (N°) = 1.99 cm, columna 2 (Combinación de Carga) = 14 cm
        ancho_col1 = 1.99
        ancho_col2 = 14.0
        alto_fila = 0.55

        for header, lista in comb_por_tipo.items():
            titulo = titulos.get(header, {}).get("en" if idioma in ["ingles", "en"] else "es", header)
            nuevo.add_paragraph(titulo)
            p = nuevo.paragraphs[-1]
            run = p.runs[0] if p.runs else p.add_run()
            run.font.name = "Arial"
            run.font.size = Pt(12)
            run.bold = True
            p.alignment = 0
            nuevo.add_paragraph()

            if idioma in ["ingles", "en"]:
                headers_tabla = ["No.", "Load Combination"]
            else:
                headers_tabla = ["N°", "Combinación de Carga"]
            tabla = nuevo.add_table(rows=1, cols=2)
            set_repeat_table_header(tabla.rows[0])
            tabla.autofit = False
            tabla.alignment = 1
            tabla.columns[0].width = Cm(ancho_col1)
            tabla.columns[1].width = Cm(ancho_col2)
            for i, h in enumerate(headers_tabla):
                cell = tabla.cell(0, i)
                cell.text = h
                p = cell.paragraphs[0]
                p.alignment = 1
                p.runs[0].font.bold = True
                p.runs[0].font.size = Pt(12)
                p.runs[0].font.name = "Arial"
                shading_elm = parse_xml(
                    r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="D9D9D9"/>'
                )
                cell._tc.get_or_add_tcPr().append(shading_elm)
                cell.width = Cm(ancho_col1 if i == 0 else ancho_col2)
            tabla.rows[0].height = Cm(alto_fila)
            tabla.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            for num, nombre in lista:
                row = tabla.add_row().cells
                row[0].text = format_number(num)
                row[1].text = nombre
                for i in range(2):
                    p = row[i].paragraphs[0]
                    p.alignment = 1
                    p.runs[0].font.size = Pt(12)
                    p.runs[0].font.name = "Arial"
                    row[i].width = Cm(ancho_col1 if i == 0 else ancho_col2)
                tabla.rows[-1].height = Cm(alto_fila)
                tabla.rows[-1].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            poner_bordes_tabla(tabla)
            nuevo.add_paragraph()
    else:
        if idioma in ["ingles", "en"]:
            nuevo.add_paragraph("Load combinations not available.")
        else:
            nuevo.add_paragraph("Combinaciones de carga no disponibles.")
    # Salto de página al final de la sección
    nuevo.add_page_break()