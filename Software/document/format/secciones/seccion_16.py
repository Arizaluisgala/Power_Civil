from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def agregar_imagenes_cargas_pp_vz(nuevo, image_slots, idioma):
    """
    Agrega imágenes de cargas y sus títulos al documento 'nuevo' según los slots de imágenes y el idioma.
    """
    if not image_slots:
        print("⚠️ No hay image_slots disponibles")
        return nuevo

    # Títulos y descripciones por idioma, usando números como clave
    titulos = {
        8:  {"es": "Peso propio de la estructura (D)", "en": "Self-weight of the structure (D)"},
        9:  {"es": "Carga de operación (Do)", "en": "Operation load (Do)"},
        10: {"es": "Carga de prueba hidrostática (F)", "en": "Hydrostatic test load (F)"},
        11: {"es": "Carga viva de la estructura (L)", "en": "Live load of the structure (L)"},
        12: {"es": "Carga temperatura (TST)", "en": "Temperature load (TST)"},
        13: {"es": "Fuerzas de fricción (TT)", "en": "Friction forces (TT)"},
        14: {"es": "Fuerzas de anclaje (TS)", "en": "Anchorage forces (TS)"},
        15: {"es": "Fuerzas de viento en dirección X (Wx)", "en": "Wind forces in X direction (Wx)"},
        16: {"es": "Fuerzas de viento en dirección -X (Wx)", "en": "Wind forces in -X direction (Wx)"},
        17: {"es": "Fuerzas de viento en dirección Z (Wz)", "en": "Wind forces in Z direction (Wz)"},
        18: {"es": "Fuerzas de viento en dirección -Z (Wz)", "en": "Wind forces in -Z direction (Wz)"},
    }
    orden = [8, 9, 10, 11, 12, 13, 14, 15, 16, 17, 18]
    lang = "en" if idioma in ["ingles", "en"] else "es"

    # Título principal
    titulo_principal = "Cargas de diseño" if lang == "es" else "Design loads"
    p_titulo = nuevo.add_paragraph(titulo_principal)
    p_titulo.alignment = 1  # Centrado
    r_titulo = p_titulo.runs[0]
    r_titulo.font.name = "Arial"
    r_titulo.font.size = Pt(12)
    r_titulo.bold = True
    r_titulo.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Alineación justificada
    nuevo.add_paragraph()  # Espacio extra después del título
    nuevo.add_paragraph()

    # Párrafo introductorio
    parrafo_intro = (
        "A continuación, se muestran las cargas aplicadas en la estructura"
        if lang == "es"
        else "Below are the loads applied to the structure"
    )
    p_intro = nuevo.add_paragraph(parrafo_intro)
    p_intro.alignment = 1
    r_intro = p_intro.runs[0]
    r_intro.font.name = "Arial"
    r_intro.font.size = Pt(12)
    r_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Alineación justificada
    nuevo.add_paragraph()  # Espacio extra después del párrafo introductorio



    # Espacio extra
    nuevo.add_paragraph("")

    # Agregar imágenes
    idx = 8
    for num_fig in orden:
        if num_fig in image_slots:
            if isinstance(image_slots[num_fig], tuple):
                ruta, _ = image_slots[num_fig]
            else:
                ruta = image_slots[num_fig]
            try:
                import os
                if not os.path.exists(ruta):
                    print(f"❌ Archivo no encontrado: {ruta}")
                    continue
                # Imagen
                p = nuevo.add_paragraph()
                p.alignment = 1
                run = p.add_run()
                run.add_picture(ruta, width=Cm(17), height=Cm(16))
                # Línea negra por todos los bordes de las imágenes
                nuevo.add_paragraph()  # Espacio extra después de la imagen
                # Título de la figura
                titulo = titulos[num_fig][lang]
                figure_word = "Figure" if lang == "en" else "Figura"
                p2 = nuevo.add_paragraph(f"{figure_word} {idx}. {titulo}")
                p2.alignment = 1
                r2 = p2.runs[0]
                r2.font.name = "Arial"
                r2.font.size = Pt(12)
                r2.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alineación centrada
                idx += 1
                # Salto de página después de cada imagen menos la última
                if idx <= len(orden):
                    nuevo.add_page_break()

            except Exception as e:
                print(f"❌ Error al agregar imagen {num_fig}: {e}")
                continue

    print(f"✅ Sección de cargas completada con {idx-1} imágenes")
    return nuevo