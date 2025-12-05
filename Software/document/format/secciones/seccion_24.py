from docx.shared import Pt, Cm
from docx.oxml import parse_xml
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
import os
import sys
project_root = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..', '..'))
sys.path.append(project_root)
sys.path.append(os.path.join(project_root, 'scripts'))

# Usar la nueva función que retorna las tablas bien separadas
from scripts.get_cortante_basal import extraer_tablas_cb_especifico_dict
# Si tienes un módulo para bordes de tabla, descomenta la siguiente línea:
from componentes.poner_bordes_tablas import poner_bordes_tabla
from componentes.set_repeat_header import set_repeat_table_header
from document.format.helpers import format_number

def seccion_cortante_basal(doc, idioma, ruta_excel):
    # Diccionario de traducción de encabezados
    traduccion_headers = {
        "Base Shear": {"es": "Cortante Basal", "en": "Base Shear"},
        "Story": {"es": "Nivel", "en": "Story"},
        "Height": {"es": "Altura", "en": "Height"},
        "Weight": {"es": "Peso", "en": "Weight"},
        "Shear": {"es": "Cortante", "en": "Shear"},
        "Cu (Table 12.8-1)": {"es": "Cu (Tabla 12.8-1)", "en": "Cu (Table 12.8-1)"},
        "Factor": {"es": "Factor", "en": "Factor"},
        "Scale": {"es": "Escalamiento", "en": "Scale"},
        "Parameter": {"es": "Parámetro", "en": "Parameter"},
        "Value": {"es": "Valor", "en": "Value"},
        # Agrega más según tus encabezados reales
    }
    """
    Inserta la sección de cortante basal en el documento doc, usando el idioma ('es' o 'en')
    y la ruta al archivo Excel de donde extraer los datos.
    """
    key = 'en' if idioma.lower() in ('ingles', 'en') else 'es'
    textos = {
        'parrafo': {
            'es': "Realizada la primera corrida, se procede a la corrección del sismo dinámico de acuerdo con lo indicado en la norma ASCE-7-XX y al chequeo de los desplazamientos en la estructura.",
            'en': "After the first run, the dynamic earthquake is corrected according to ASCE-7-XX and the displacements in the structure are checked."
        },
        'titulo': {
            'es': "Cortante basal y escalamiento del sismo",
            'en': "Base shear and earthquake scaling"
        },
        'intro': {
            'en': "The following tables summarize the main parameters and results of the base shear calculation according to the project criteria and seismic code.",
            'es': "Las siguientes tablas resumen los principales parámetros y resultados del cálculo de cortante basal según los criterios del proyecto y el código sísmico."
        },
        'sin_datos': {
            'en': "No data available for this section.",
            'es': "No hay datos disponibles para esta sección."
        }
    }

    try:
        # Extraer las tablas bien separadas
        tablas_cb = extraer_tablas_cb_especifico_dict(ruta_excel)
    except ValueError as e:
        print(f"[WARNING] No se pudo generar la sección de cortante basal: {e}")
        doc.add_paragraph(textos['sin_datos'][key])
        return

    # Párrafo inicial justificado
    doc.add_paragraph()
    p = doc.add_paragraph(textos['parrafo'][key])
    p.runs[0].font.name = "Arial"; p.runs[0].font.size = Pt(12)
    p.alignment = 3  # Justificado
    doc.add_paragraph()
    # Título principal alineado a la izquierda
    p = doc.add_paragraph(textos['titulo'][key])
    r = p.runs[0]; r.font.name = "Arial"; r.font.size = Pt(12); r.bold = True
    p.alignment = 0  # Izquierda
    doc.add_paragraph()

    # Crear cada tabla
    for idx, tabla_dict in enumerate(tablas_cb):
        headers = tabla_dict['headers']
        rows = tabla_dict['rows']
        if not rows:
            continue
        # Para la primera tabla (solo mostrar Cu (Table 12.8-1) y su valor, celda gris, sin fila extra)
        if idx == 0:
            # Buscar la fila con valor en la columna de encabezado que contenga 'Cu (Table 12.8-1)'
            cu_val = None
            for row in rows:
                for i, h in enumerate(headers):
                    if 'Cu (Table 12.8-1)' in str(row[i]):
                        cu_val = row[i+1] if i+1 < len(row) else ''
                        break
                if cu_val is not None:
                    break
            if cu_val is not None:
                tabla = doc.add_table(rows=1, cols=2)
                tabla.autofit = False
                tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
                tabla.style = 'Table Grid'
                # Primera celda: encabezado gris
                cell0 = tabla.cell(0, 0)
                cell0.text = 'Cu (Table 12.8-1)'
                p0 = cell0.paragraphs[0]; p0.alignment = 1
                while len(p0.runs) > 1:
                    p0.runs[-1].clear(); del p0.runs[-1]
                run0 = p0.runs[0]
                run0.text = 'Cu (Table 12.8-1)'; run0.font.name = "Arial"; run0.font.size = Pt(11); run0.bold = True
                rPr = run0._element.get_or_add_rPr()
                sz = parse_xml(r'<w:sz xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="22"/>')
                szCs = parse_xml(r'<w:szCs xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="22"/>')
                rPr.append(sz); rPr.append(szCs)
                shade = parse_xml(r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="D9D9D9"/>')
                cell0._tc.get_or_add_tcPr().append(shade)
                cell0.width = Cm(4.5)
                cell0.vertical_alignment = 1
                # Segunda celda: valor
                cell1 = tabla.cell(0, 1)
                # Formatear a 2 decimales si es número
                try:
                    cu_val_float = float(cu_val)
                    cell1.text = f"{cu_val_float:.2f}"
                except Exception:
                    cell1.text = str(cu_val)
                p1 = cell1.paragraphs[0]; p1.alignment = 1
                while len(p1.runs) > 1:
                    p1.runs[-1].clear(); del p1.runs[-1]
                run1 = p1.runs[0]
                run1.text = cell1.text; run1.font.name = "Arial"; run1.font.size = Pt(11)
                cell1.width = Cm(2.5)
                cell1.vertical_alignment = 1
                tabla.rows[0].height = Cm(1.0)
                tabla.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
                try:
                    from componentes.poner_bordes_tablas import poner_bordes_tabla
                    poner_bordes_tabla(tabla)
                except Exception:
                    pass
                doc.add_paragraph()
            continue
        # El resto de tablas sí llevan encabezado y estilo igual a otras secciones
        tabla = doc.add_table(rows=1, cols=len(headers))
        set_repeat_table_header(tabla.rows[0])
        tabla.autofit = False
        tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
        tabla.style = 'Table Grid'
        for j, h in enumerate(headers):
            cell = tabla.cell(0, j)
            # Traducir encabezado según idioma
            h_str = str(h)
            h_trad = traduccion_headers.get(h_str, {}).get(key, h_str)
            if len(h_trad) > 22 and ' ' in h_trad:
                parts = h_trad.split(' ')
                mid = len(parts) // 2
                h_trad = ' '.join(parts[:mid]) + '\n' + ' '.join(parts[mid:])
            cell.text = h_trad
            p0 = cell.paragraphs[0]; p0.alignment = 1
            while len(p0.runs) > 1:
                p0.runs[-1].clear(); del p0.runs[-1]
            run = p0.runs[0]
            run.text = h_trad; run.font.name = "Arial"; run.font.size = Pt(11); run.bold = True
            rPr = run._element.get_or_add_rPr()
            sz = parse_xml(r'<w:sz xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="22"/>')
            szCs = parse_xml(r'<w:szCs xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="22"/>')
            rPr.append(sz); rPr.append(szCs)
            shade = parse_xml(r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="D9D9D9"/>')
            cell._tc.get_or_add_tcPr().append(shade)
            cell.width = Cm(3.2)
            cell.vertical_alignment = 1
        tabla.rows[0].height = Cm(1.0)
        tabla.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        for row in rows:
            cells = tabla.add_row().cells
            for j, val in enumerate(row):
                try:
                    cells[j].text = format_number(float(val))
                except (ValueError, TypeError):
                    cells[j].text = str(val) if val is not None else ""
                for paragraph in cells[j].paragraphs:
                    paragraph.alignment = 1
                    for run in paragraph.runs:
                        run.font.name = "Arial"
                        run.font.size = Pt(10)
                cells[j].width = Cm(3.2)
                cells[j].vertical_alignment = 1
        # Bordes y espacio
        try:
            from componentes.poner_bordes_tablas import poner_bordes_tabla
            poner_bordes_tabla(tabla)
        except Exception:
            pass
        doc.add_paragraph()

    doc.add_paragraph()