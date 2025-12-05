import sys
import os
from docx.shared import Pt, Cm
from docx.oxml import parse_xml
from docx.enum.table import WD_ROW_HEIGHT_RULE
import pandas as pd
project_root = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..', '..'))
sys.path.append(project_root)
sys.path.append(os.path.join(project_root, 'scripts'))
sys.path.append(os.path.join(project_root, 'componentes'))
from scripts.extract_tables_of_excel_load import extract_target_tables
from componentes.poner_bordes_tablas import poner_bordes_tabla
from componentes.set_repeat_header import set_repeat_table_header
from document.format.helpers import format_number


def merge_vertical_cells(table, col_idx):
    """Fusiona celdas verticalmente en la columna col_idx si tienen el mismo valor o están vacías."""
    start = 1  # Empieza después del encabezado
    while start < len(table.rows):
        current_cell = table.cell(start, col_idx)
        current_text = current_cell.text.strip()
        end = start
        # Fusionar celdas vacías o iguales
        while end + 1 < len(table.rows):
            next_cell = table.cell(end + 1, col_idx)
            next_text = next_cell.text.strip()
            if next_text == "" or next_text == current_text:
                end += 1
            else:
                break
        if end > start:
            current_cell.merge(table.cell(end, col_idx))
        start = end + 1


def cargas_aplicadas(nuevo, idioma, excel_cargas_path):
    tablas_cargas = extract_target_tables(excel_cargas_path)

    # Orden y títulos fijos - incluir todos los tipos de tabla
    ORDEN = [6, 5, 3, 2, 0, 4, 1]  # Wind Definitions, SelfWeights, Beam Loads, Temperature Loads, One-way floor, Nodal Loads, Response Spectrum
    TITULOS_TABLAS = {
        0: {"es": "Cargas de piso unidireccionales", "en": "One-Way Floor Loads"},
        1: {"es": "Cargas de espectro de respuesta", "en": "Response Spectrum Loads"},
        2: {"es": "Cargas de temperatura", "en": "Temperature Loads"},
        3: {"es": "Cargas de viga", "en": "Beam Loads"},
        4: {"es": "Cargas nodales", "en": "Nodal Loads"},
        5: {"es": "Pesos propios", "en": "SelfWeights"},
        6: {"es": "Definiciones de carga de viento", "en": "Wind Load Definitions"},
    }
    # Traducción de encabezados
    TRADUCCION_HEADERS = {
        "L/C": {"es": "Caso de Carga", "en": "L/C"},
        "Type": {"es": "Tipo", "en": "Type"},
        "Title": {"es": "Título", "en": "Title"},
        "Beam": {"es": "Viga", "en": "Beam"},
        "Load": {"es": "Carga", "en": "Load"},
        "Node": {"es": "Nodo", "en": "Node"},
        "FX": {"es": "FX", "en": "FX"},
        "FY": {"es": "FY", "en": "FY"},
        "FZ": {"es": "FZ", "en": "FZ"},
        "Temperature": {"es": "Temperatura", "en": "Temperature"},
        "Direction": {"es": "Dirección", "en": "Direction"},
        "Intensity": {"es": "Intensidad", "en": "Intensity"},
        "Height": {"es": "Altura", "en": "Height"},
    }

    if not tablas_cargas:
        if idioma in ["ingles", "en"]:
            nuevo.add_paragraph("No applied load tables found.")
        else:
            nuevo.add_paragraph("No se encontraron tablas de cargas aplicadas.")
        return

    # Encabezados esperados para cada tipo de tabla (ajustados a todos los patrones según TARGET_PATTERNS)
    HEADERS_ESPERADOS = {
        0: [  # One-way floor loads
            ["L/C", "Direction", "Load", "Min X", "Max X", "Min Y", "Max Y", "Min Z", "Max Z"],
        ],
        1: [  # Response Spectrum Loads
            ["L/C", "Parameter", "Value"],
        ],
        2: [  # Temperature Loads
            ["L/C", "Beam", "Type", "Axial", "Cross", "Side", "Elongation", "Strain Rate"],
        ],
        3: [  # Beam Loads
            ["L/C", "Beam", "Type", "Direction", "Fa", "Da", "Fb", "Db", "Ecc."],
        ],
        4: [  # Nodal Loads
            ["L/C", "Node", "FX", "FY", "FZ", "MX", "MY", "MZ", "Type"],
        ],
        5: [  # SelfWeights
            ["L/C", "Direction", "Factor", "Assigned Geometry"],
        ],
        6: [  # Wind Load Definitions
            ["Type", "Title", "Intensity", "Height"],
        ],
        # Patrones adicionales que pueden aparecer
        7: [  # Wind Loads (aparece en el documento)
            ["L/C", "Direction", "Type", "Factor"],
        ],
    }

    # Función para comparar encabezados (ignora espacios extra y mayúsculas/minúsculas)
    def headers_match(cols, expected_list):
        cols_norm = [str(c).strip().lower() for c in cols]
        for expected in expected_list:
            expected_norm = [str(e).strip().lower() for e in expected]
            if cols_norm == expected_norm:
                return True
        return False

    # Filtrar y ordenar las tablas según ORDEN y encabezado esperado
    tablas_filtradas = []
    
    # Primero, intentar coincidencias exactas con los patrones esperados
    for idx in ORDEN:
        for t in tablas_cargas:
            if t["pattern_index"] == idx:
                df = t["data"].copy()
                df = df.dropna(how='all')
                # Si la primera fila parece de unidades, ignórala para comparar encabezados
                if (
                    df.shape[0] > 0
                    and all(isinstance(val, str) and (val.strip() == "" or "(" in val or "/" in val or "°" in val) for val in df.iloc[0])
                ):
                    cols_to_check = list(df.columns)
                else:
                    cols_to_check = list(df.columns)
                
                # Verificar si coincide con algún patrón esperado
                if headers_match(cols_to_check, HEADERS_ESPERADOS.get(idx, [])):
                    tablas_filtradas.append(t)
                    
    # Si no encontramos suficientes tablas, agregar las que no coincidieron exactamente
    if len(tablas_filtradas) < 4:  # Esperamos al menos las principales tablas
        for t in tablas_cargas:
            if t not in tablas_filtradas:
                df = t["data"].copy()
                df = df.dropna(how='all')
                if len(df) > 0:  # Si tiene datos, agregarla
                    tablas_filtradas.append(t)

    for tabla in tablas_filtradas:
        pattern_idx = tabla["pattern_index"]
        
        # Usar el título del diccionario si existe, sino generar uno basado en las columnas
        if pattern_idx in TITULOS_TABLAS:
            titulo_tabla = TITULOS_TABLAS[pattern_idx]["en"] if idioma in ["ingles", "en"] else TITULOS_TABLAS[pattern_idx]["es"]
        else:
            # Generar título basado en las columnas de la tabla
            df_temp = tabla["data"].copy()
            cols = list(df_temp.columns)
            if "Wind" in str(cols) or "Type" in cols and "Title" in cols:
                titulo_tabla = "Wind Load Definitions" if idioma in ["ingles", "en"] else "Definiciones de carga de viento"
            elif "Temperature" in str(cols) or "Beam" in cols and "Axial" in cols:
                titulo_tabla = "Temperature Loads" if idioma in ["ingles", "en"] else "Cargas de temperatura"
            elif "Direction" in cols and "Factor" in cols and "L/C" in cols:
                titulo_tabla = "Wind Loads" if idioma in ["ingles", "en"] else "Cargas de viento"
            else:
                titulo_tabla = f"Applied Loads {pattern_idx}" if idioma in ["ingles", "en"] else f"Cargas Aplicadas {pattern_idx}"
        p_titulo = nuevo.add_paragraph(titulo_tabla)
        run_titulo = p_titulo.runs[0] if p_titulo.runs else p_titulo.add_run()
        run_titulo.font.name = "Arial"
        run_titulo.font.size = Pt(12)
        run_titulo.bold = True
        p_titulo.alignment = 0  # Izquierda
        nuevo.add_paragraph()  # Espacio

        df = tabla["data"].copy()
        df = df.dropna(how='all')
        for col in df.columns:
            df[col] = df[col].apply(lambda x: x if pd.notna(x) and str(x).strip() != '' else pd.NA)
        df = df.dropna(how='all')
        if len(df) == 0:
            continue

        # --- COMBINAR ENCABEZADO Y UNIDADES EN UNA SOLA FILA ---
        if (
            df.shape[0] > 0
            and all(isinstance(val, str) and (val.strip() == "" or "(" in val or "/" in val or "°" in val) for val in df.iloc[0])
        ):
            unidades = list(df.iloc[0])
            nombres = list(df.columns)
            nuevos_headers = []
            for nombre, unidad in zip(nombres, unidades):
                nombre_completo = nombre
                # Corrección de encabezados cortados
                if nombre.lower().startswith("direcci"):
                    nombre_completo = "Dirección"
                if unidad and unidad.strip():
                    nombre_completo = f"{nombre_completo.strip()} {unidad.strip()}"
                nuevos_headers.append(nombre_completo)
            df.columns = nuevos_headers
            df = df.iloc[1:].reset_index(drop=True)  # Eliminar la fila de unidades de los datos

        # Traducir encabezados si es necesario
        if idioma in ["ingles", "en"]:
            # Aquí puedes agregar traducción si lo deseas
            pass

        ncols = len(df.columns)
        nrows = len(df)
        # Calcular anchos de columna según el largo del encabezado
        max_len = [max(len(str(col)), max((len(str(val)) for val in df[col] if pd.notna(val)), default=0)) for col in df.columns]
        total_len = sum(max_len)
        ancho_total = 16  # Puedes ajustar este valor según tu documento
        col_widths = [ancho_total * l / total_len for l in max_len]
        
        # Ajuste especial: Si hay una columna "Direction" en cargas de piso, darle ancho fijo de 2.4 cm
        for j, col in enumerate(df.columns):
            if "Direction" in str(col) and pattern_idx == 0:  # Solo para One-way floor loads (índice 0)
                col_widths[j] = 2.4
                # Recalcular el resto proporcionalmente
                remaining_width = ancho_total - 2.4
                remaining_cols = [i for i in range(len(col_widths)) if i != j]
                if remaining_cols:
                    remaining_max_len = [max_len[i] for i in remaining_cols]
                    remaining_total = sum(remaining_max_len)
                    if remaining_total > 0:
                        for i in remaining_cols:
                            col_widths[i] = remaining_width * max_len[i] / remaining_total

        tabla_word = nuevo.add_table(rows=nrows+1, cols=ncols)
        set_repeat_table_header(tabla_word.rows[0])
        if pattern_idx not in [1, 5]:
            set_repeat_table_header(tabla_word.rows[1])
        tabla_word.autofit = False

        # Encabezados traducidos y con tamaño 9
        for j, col in enumerate(df.columns):
            cell = tabla_word.cell(0, j)
            # Traducir encabezado según idioma
            if idioma in ["ingles", "en"]:
                col_trad = TRADUCCION_HEADERS.get(str(col), {}).get("en", str(col))
            else:
                col_trad = TRADUCCION_HEADERS.get(str(col), {}).get("es", str(col))
            cell.text = str(col_trad)
            p = cell.paragraphs[0]
            p.alignment = 1
            run = p.runs[0] if p.runs else p.add_run()
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.name = "Arial"
            shading_elm = parse_xml(
                r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="D9D9D9"/>'
            )
            cell._tc.get_or_add_tcPr().append(shading_elm)
            tabla_word.columns[j].width = Cm(col_widths[j])

        # Alto encabezado: permite crecer si el texto es largo
        tabla_word.rows[0].height = Cm(0.55)
        tabla_word.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

        # Datos con letra tamaño 11
        for i, (_, row) in enumerate(df.iterrows()):
            for j, val in enumerate(row):
                cell = tabla_word.cell(i+1, j)
                if pd.isna(val):
                    cell.text = ""
                else:
                    cell.text = format_number(val)
                p = cell.paragraphs[0]
                p.alignment = 1
                run = p.runs[0] if p.runs else p.add_run()
                run.font.size = Pt(11)
                run.font.name = "Arial"
                # Segunda fila (i==0) debajo del encabezado - aplicar formato especial
                if i == 0 and pattern_idx not in [1, 5]:
                    shading_elm = parse_xml(
                        r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="D9D9D9"/>'
                    )
                    cell._tc.get_or_add_tcPr().append(shading_elm)
                    run.font.bold = True
            tabla_word.rows[i+1].height = Cm(0.55)
            tabla_word.rows[i+1].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

        # Fusionar verticalmente la columna L/C o la primera columna si existe
        merge_col_name = "L/C" if "L/C" in df.columns else df.columns[0]
        merge_col_idx = list(df.columns).index(merge_col_name)
        merge_vertical_cells(tabla_word, merge_col_idx)

        poner_bordes_tabla(tabla_word)
        nuevo.add_paragraph()  # Espacio entre tablas
    
 