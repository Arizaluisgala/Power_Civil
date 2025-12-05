from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def add_mass_section(nuevo,image_slots, idioma):
    if idioma== 'en':
        texto = (
            'To perform the seismic analysis, a reference load of type "mass" is previously generated, '
            'which includes all the masses to be considered.'
        )
        figura = "Figure 7. Reference load type mass"
    else:
        texto = (
            'Para realizar el análisis sísmico se genera previamente una carga de referencia tipo “mass”, '
            'en la que se incluyen todas las masas a considerar.'
        )
        figura = "Figura 7. Carga de referencia tipo mass"
        
        
    if not image_slots:
        print("⚠️ No hay image_slots disponibles")
        return nuevo    

    # Párrafo centralizado con texto
    p = nuevo.add_paragraph()
    run = p.add_run(texto)
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    nuevo.add_paragraph("")  # Espacio extra

    # Imagen centrada, tamaño fijo
    image_path = image_slots.get(8, None)
    p_img = nuevo.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_img = p_img.add_run()
    run_img.add_picture(image_path, width=Cm(17), height=Cm(12))

    # Título de la figura debajo, centrado, Arial 12
    nuevo.add_paragraph()  # Espacio extra
    p_title = nuevo.add_paragraph(figura)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.runs[0]
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(12)

    nuevo.add_page_break()