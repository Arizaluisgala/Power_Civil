from docx.shared import Pt

def agregar_titulo_estructura(doc, estructura):
        titulo_parrafo = doc.add_paragraph()
        run = titulo_parrafo.add_run(estructura.upper())
        run.font.name = "Arial"
        run.font.size = Pt(14)
        run.bold = True
        titulo_parrafo.alignment = 0  # Izquierda
        doc.add_paragraph()  # Agregar un salto de línea después del título
        return estructura.upper()