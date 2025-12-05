import os
from docx.shared import Inches
from docx.oxml import parse_xml

def clonar_encabezado(base_header, nuevo_header, reemplazos, logo_path=None):
    for elem in list(nuevo_header._element.xpath('./w:p | ./w:tbl')):
        nuevo_header._element.remove(elem)
    for elem in base_header._element.xpath('./w:p | ./w:tbl'):
        nuevo = parse_xml(elem.xml)
        for node in nuevo.xpath('.//w:t'):
            txt = node.text or ""
            if txt in reemplazos:
                node.text = reemplazos[txt]
        nuevo_header._element.append(nuevo)
    if logo_path and os.path.isfile(logo_path):
        tablas = nuevo_header.tables
        if tablas:
            celda_logo = tablas[0].cell(0, 0)
            par = celda_logo.paragraphs[0]
            for run in par.runs:
                p_element = par._element
                p_element.remove(run._element)
            run = par.add_run()
            run.add_picture(logo_path, width=Inches(1.5))
