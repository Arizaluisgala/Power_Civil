from docx.shared import Cm, Pt

def insertar_imagen_con_leyenda(doc, ruta, leyenda, ancho_cm=16.51):
        p = doc.add_paragraph()
        run = p.add_run()
        run.add_picture(ruta, width=Cm(ancho_cm))
        p2 = doc.add_paragraph(leyenda)
        p2.alignment = 1  # centrar
        r2 = p2.runs[0]
        r2.font.name = "Arial"
        r2.font.size = Pt(10)
        doc.add_page_break()
