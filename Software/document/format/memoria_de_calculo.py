from docx import Document
from dotenv import load_dotenv
import importlib
import sys
import os
import datetime

# Componentes
from componentes.copiar_configuracion_pagina import copiar_configuracion_pagina
from componentes.clonar_encabezado import clonar_encabezado
from componentes.copiar_estilos_y_bordes import copiar_estilos_y_bordes
from componentes.agregar_titulo_estructura import agregar_titulo_estructura
from componentes.validar_idioma import validar_idioma
from componentes.validar_version import validar_version
from componentes.extraer_textos import extraer_textos_header
from componentes.pedir_reemplazos import pedir_reemplazos

# secciones
from secciones.seccion_1 import agregar_imagenes_seccion
from secciones.seccion_2 import tablas_miembros_materiales_soportes  
from secciones.seccion_3 import casos_de_carga_primarias
from secciones.seccion_4 import generar_espectro
from secciones.seccion_5 import cargas_aplicadas
from secciones.seccion_6 import casos_de_carga_combinados
from secciones.seccion_7 import verificacion_de_deflexion_vertical_en_vigas
from secciones.seccion_8 import verificacion_de_deflexión_horizontal_en_vigas
from secciones.seccion_9 import verificacion_de_desplazamientos_por_viento
from secciones.seccion_10 import verificacion_por_sismo
from secciones.seccion_11 import ratios_tables
from secciones.seccion_12 import leer_std_a_texto
from secciones.seccion_13 import computos_metricos
from secciones.seccion_14 import reacciones_de_la_estructura
from secciones.seccion_15 import add_mass_section
from secciones.seccion_16 import agregar_imagenes_cargas_pp_vz
from secciones.seccion_17 import viga_image_v
from secciones.seccion_18 import viga_image_h
from secciones.seccion_19 import columna_image_hx_viento
from secciones.seccion_20 import columna_image_hz_viento
from secciones.seccion_21 import columna_image_hx_sismo
from secciones.seccion_22 import columna_image_hz_sismo
from secciones.seccion_23 import image_ratio
from secciones.seccion_24 import seccion_cortante_basal
import sys
import os
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..')))

from scripts.config import IMAGE_SLOTS
from scripts.screenshots import select_region_and_save

load_dotenv()


def crear_memoria_de_calculo(
    plantilla_path,
    output_path,
    logo_path=None,
    default_logo_path=None,
    idioma=None,
    estructura=None,
    version=None,
    reemplazos=None,
    image_slots=None,
    excel_file_path=None, 
    excel_file_path_cargas=None,
    excel_file_path_sismo=None,
    tomar_imagenes="s",
    mostrar_seccion_8="s",
    mostrar_cargas="s",
    progress_callback=None,
):
    total_steps = 22 if version == "simple" else 30  # Aproximación de pasos
    current_step = 0

    def update_progress(description):
        nonlocal current_step
        current_step += 1
        if progress_callback:
            progress_callback(current_step / total_steps, description)
        print(f"[{current_step}/{total_steps}] {description}")

    try:
        from staad_automation.get_path_of_staad_connetc import get_path_of_staad_connect
        path_std = get_path_of_staad_connect()
        if path_std and os.path.isfile(path_std):
            os.chdir(os.path.dirname(path_std))
            import staad_automation.get_path_of_staad_connetc as staad_mod
            importlib.reload(staad_mod)
            update_progress(f"Directorio de trabajo cambiado a: {os.getcwd()}")
        else:
            update_progress("No se encontró archivo .std para cambiar directorio.")
    except Exception as e:
        update_progress(f"Error al forzar contexto STAAD: {e}")

    try:
        timestamp = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        directorio, nombre_archivo = os.path.split(output_path)
        nombre_base, extension = os.path.splitext(nombre_archivo)
        nuevo_nombre_archivo = f"{nombre_base}_{timestamp}{extension}"
        output_path = os.path.join(directorio, nuevo_nombre_archivo)

        update_progress("Validando parámetros de entrada...")
        if not estructura:
            raise ValueError("Debes ingresar un nombre válido para la estructura.")
        validar_idioma(idioma)
        if version:
            version = version.strip().lower()
        validar_version(version)

        update_progress("Abriendo plantilla base...")
        base = Document(plantilla_path)
        nuevo = Document()

        update_progress("Configurando estilos y formato...")
        copiar_configuracion_pagina(base.sections[0], nuevo.sections[0])
        copiar_estilos_y_bordes(base, nuevo)

        update_progress("Procesando encabezado...")
        base_header = base.sections[0].header
        textos_header = extraer_textos_header(base_header)

        if not reemplazos:
            reemplazos = pedir_reemplazos(textos_header)
        else:
            for txt in textos_header:
                if txt not in reemplazos:
                    reemplazos[txt] = txt

        update_progress("Aplicando datos del proyecto...")
        if not logo_path or not os.path.exists(logo_path):
            if default_logo_path and os.path.exists(default_logo_path):
                logo_path = default_logo_path
            else:
                logo_path = None
        
        tiene_logo = False
        try:
            tablas = nuevo.sections[0].header.tables
            if tablas:
                celda_logo = tablas[0].cell(0, 0)
                for par in celda_logo.paragraphs:
                    for run in par.runs:
                        if run._element.xpath('.//w:drawing'):
                            tiene_logo = True
                            break
                    if tiene_logo:
                        break
        except Exception:
            pass
        
        clonar_encabezado(
            base_header,
            nuevo.sections[0].header,
            reemplazos,
            logo_path if not tiene_logo else None
        )

        if "TÍTULO DE LA ESTRUCTURA" not in textos_header:
            agregar_titulo_estructura(nuevo, estructura)

        if version == "simple":
            # Simple version implementation
            if tomar_imagenes == "s" or tomar_imagenes == "y":
                update_progress("Insertando imágenes del proyecto...")
                agregar_imagenes_seccion(nuevo, image_slots, idioma)

            if excel_file_path and os.path.exists(excel_file_path):
                update_progress("Procesando tablas de datos técnicos...")
                tablas_miembros_materiales_soportes(nuevo, idioma, excel_file_path)
            
            if excel_file_path and os.path.exists(excel_file_path):
                update_progress("Generando casos de carga primarios...")
                casos_de_carga_primarias(nuevo,idioma)

            if excel_file_path_cargas and os.path.exists(excel_file_path_cargas):
                update_progress("Generando espectros de diseño...")
                generar_espectro(nuevo, idioma, excel_file_path_cargas)

            if mostrar_cargas and (str(mostrar_cargas).lower() in ["s", "si", "true", "1"]) and excel_file_path_cargas and os.path.exists(excel_file_path_cargas):
                update_progress("Procesando cargas aplicadas...")
                cargas_aplicadas(nuevo,idioma,excel_file_path_cargas)

            if excel_file_path and os.path.exists(excel_file_path):
                update_progress("Generando casos de carga combinados...")
                casos_de_carga_combinados(nuevo,idioma)

            archivo_sismo = excel_file_path_sismo if excel_file_path_sismo and os.path.exists(excel_file_path_sismo) else excel_file_path_cargas
            if archivo_sismo and os.path.exists(archivo_sismo):
                update_progress("Procesando cortante basal...")
                seccion_cortante_basal(nuevo,idioma,archivo_sismo)

            if excel_file_path and os.path.exists(excel_file_path):
                update_progress("Verificando deflexión vertical en vigas...")
                verificacion_de_deflexion_vertical_en_vigas(nuevo,idioma,excel_file_path)

            if mostrar_seccion_8 and (str(mostrar_seccion_8).lower() in ["s", "si", "true", "1"]):
                update_progress("Verificando deflexión horizontal en vigas...")
                verificacion_de_deflexión_horizontal_en_vigas(nuevo,idioma,excel_file_path)

            if excel_file_path and os.path.exists(excel_file_path):
                update_progress("Analizando desplazamientos por viento...")
                verificacion_de_desplazamientos_por_viento(nuevo,idioma,excel_file_path)

            if excel_file_path and os.path.exists(excel_file_path):
                update_progress("Realizando verificación sísmica...")
                verificacion_por_sismo(nuevo,idioma,excel_file_path)

            if excel_file_path and os.path.exists(excel_file_path):
                update_progress("Calculando ratios de diseño...")
                ratios_tables(nuevo,idioma,excel_file_path)

            if excel_file_path and os.path.exists(excel_file_path):
                update_progress("Generando cómputos métricos...")
                computos_metricos(nuevo,idioma,excel_file_path)

            if excel_file_path and os.path.exists(excel_file_path):
                update_progress("Calculando reacciones de la estructura...")
                reacciones_de_la_estructura(nuevo,idioma,excel_file_path)
            
            update_progress("Procesando archivo de modelo STAAD...")
            leer_std_a_texto(nuevo,idioma)

        else: # Version Completa
            # Full version implementation
            if tomar_imagenes == "s" or tomar_imagenes == "y":
                update_progress("Insertando imágenes del proyecto...")
                agregar_imagenes_seccion(nuevo, image_slots, idioma)

            if excel_file_path and os.path.exists(excel_file_path):
                update_progress("Procesando tablas de datos técnicos...")
                tablas_miembros_materiales_soportes(nuevo, idioma, excel_file_path)

            if excel_file_path and os.path.exists(excel_file_path):
                update_progress("Generando casos de carga primarios...")
                casos_de_carga_primarias(nuevo,idioma)
            
            if tomar_imagenes == "s" or tomar_imagenes == "y":
                update_progress("Insertando imágenes de masa...")
                add_mass_section(nuevo, image_slots, idioma)

            if excel_file_path_cargas and os.path.exists(excel_file_path_cargas):
                update_progress("Generando espectros de diseño...")
                generar_espectro(nuevo, idioma, excel_file_path_cargas)

            if tomar_imagenes == "s" or tomar_imagenes == "y":
                update_progress("Insertando imágenes de cargas...")
                agregar_imagenes_cargas_pp_vz(nuevo, image_slots, idioma)

            if mostrar_cargas and (str(mostrar_cargas).lower() in ["s", "si", "true", "1"]) and excel_file_path_cargas and os.path.exists(excel_file_path_cargas):
                update_progress("Procesando cargas aplicadas...")
                cargas_aplicadas(nuevo,idioma,excel_file_path_cargas)

            if excel_file_path and os.path.exists(excel_file_path):
                update_progress("Generando casos de carga combinados...")
                casos_de_carga_combinados(nuevo,idioma)

            archivo_sismo = excel_file_path_sismo if excel_file_path_sismo and os.path.exists(excel_file_path_sismo) else excel_file_path_cargas
            if archivo_sismo and os.path.exists(archivo_sismo):
                update_progress("Procesando cortante basal...")
                seccion_cortante_basal(nuevo,idioma,archivo_sismo)

            if excel_file_path and os.path.exists(excel_file_path):
                update_progress("Verificando deflexión vertical en vigas...")
                verificacion_de_deflexion_vertical_en_vigas(nuevo,idioma,excel_file_path)

            if tomar_imagenes == "s" or tomar_imagenes == "y":
                update_progress("Insertando imágenes de deflexión vertical...")
                viga_image_v(nuevo, image_slots, idioma)
                update_progress("Insertando imágenes de deflexión horizontal...")
                viga_image_h(nuevo, image_slots, idioma)

            if mostrar_seccion_8 and (str(mostrar_seccion_8).lower() in ["s", "si", "true", "1"]):
                update_progress("Verificando deflexión horizontal en vigas...")
                verificacion_de_deflexión_horizontal_en_vigas(nuevo,idioma,excel_file_path)

            if excel_file_path and os.path.exists(excel_file_path):
                update_progress("Analizando desplazamientos por viento...")
                verificacion_de_desplazamientos_por_viento(nuevo,idioma,excel_file_path)

            if tomar_imagenes == "s" or tomar_imagenes == "y":
                update_progress("Insertando imágenes de desplazamiento por viento (X)...")
                columna_image_hx_viento(nuevo, image_slots, idioma)
                update_progress("Insertando imágenes de desplazamiento por viento (Z)...")
                columna_image_hz_viento(nuevo, image_slots, idioma)

            if excel_file_path and os.path.exists(excel_file_path):
                update_progress("Realizando verificación sísmica...")
                verificacion_por_sismo(nuevo,idioma,excel_file_path)

            if tomar_imagenes == "s" or tomar_imagenes == "y":
                update_progress("Insertando imágenes de desplazamiento sísmico (X)...")
                columna_image_hx_sismo(nuevo, image_slots, idioma)
                update_progress("Insertando imágenes de desplazamiento sísmico (Z)...")
                columna_image_hz_sismo(nuevo, image_slots, idioma)

            if excel_file_path and os.path.exists(excel_file_path):
                update_progress("Calculando ratios de diseño...")
                ratios_tables(nuevo,idioma,excel_file_path)

            if tomar_imagenes == "s" or tomar_imagenes == "y":
                update_progress("Insertando imágenes de ratios...")
                image_ratio(nuevo, image_slots, idioma)

            if excel_file_path and os.path.exists(excel_file_path):
                update_progress("Generando cómputos métricos...")
                computos_metricos(nuevo,idioma,excel_file_path)

            if excel_file_path and os.path.exists(excel_file_path):
                update_progress("Calculando reacciones de la estructura...")
                reacciones_de_la_estructura(nuevo,idioma,excel_file_path)
            
            update_progress("Procesando archivo de modelo STAAD...")
            leer_std_a_texto(nuevo,idioma)

        update_progress("Guardando documento final...")
        nuevo.save(output_path)
        update_progress(f"Documento guardado en: {output_path}")

        return {"success": True, "output_path": output_path}

    except Exception as e:
        print(f"[ERROR] Error al generar memoria: {e}")
        return {"success": False, "error": str(e)}

if __name__ == "__main__":
    # Ejemplo de uso en modo consola
    plantilla_path = os.getenv("TEMPLATE_PATH", None)
    output_path = os.getenv("OUT_DOC_PATH_M", "memoria_de_calculo.docx")
    logo_path = os.getenv("LOGO_PATH", "image.png")
    excel_file_path = os.getenv("EXCEL_TEMPLATE", None)
    excel_file_path_cargas = os.getenv("REPORT_FILE_PATH", None)
    estructura = "Test"
    idioma = "es"
    version = "completa" 
    reemplazos = None
    
    capturadas = {}
    tomar_imagenes = "n"
    
    mostrar_seccion_8 = "s"

    def console_progress(progress, description):
        print(f"Progreso: {int(progress * 100)}% - {description}")

    resultado = crear_memoria_de_calculo(
        plantilla_path,
        output_path,
        logo_path=logo_path,
        idioma=idioma,
        estructura=estructura,
        version=version,
        reemplazos=reemplazos,
        image_slots=capturadas,
        excel_file_path=excel_file_path,
        excel_file_path_cargas=excel_file_path_cargas,
        tomar_imagenes=tomar_imagenes,
        mostrar_seccion_8=mostrar_seccion_8,
        progress_callback=console_progress
    )
    if resultado and resultado.get("success"):
        print("[OK] Memoria de cálculo creada exitosamente.")
    else:
        print(f"[ERROR] Error al generar memoria: {resultado.get('error') if resultado else 'Error desconocido'}")
