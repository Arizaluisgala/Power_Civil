from docx.shared import Cm, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def agregar_borde_celda(celda, color="000000", size=4):
    """
    Agrega un borde negro delgado a una celda de tabla de docx.
    size=4 equivale a 0.5pt aprox (milimétrico).
    """
    tc = celda._tc
    tcPr = tc.get_or_add_tcPr()
    for border_name in ("top", "left", "bottom", "right"):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(size))  # 4 = 0.5pt, 8 = 1pt
        border.set(qn('w:space'), "0")
        border.set(qn('w:color'), color)
        tcPr.append(border)

def agregar_imagenes_seccion(nuevo, image_slots, idioma):
    """
    Agrega imágenes y sus títulos al documento 'nuevo' según los slots de imágenes y el idioma.
    Cada imagen tendrá un borde negro milimétrico.
    """
    if not image_slots:
        print("⚠️ No hay image_slots disponibles")
        return nuevo

    print(f"🔍 Procesando image_slots: {image_slots}")
    titulos = {
        1: {"es": "Vista isométrica 3D", "en": "3D isometric view"},
        2: {"es": "Isometría con dimensiones", "en": "Isometry with dimensions"},
        3: {"es": "Numeración de nodos", "en": "Node numbering view"},
        4: {"es": "Numeración de vigas", "en": "Beam numbering view"},
        5: {"es": "Perfiles", "en": "Profiles"},
        6: {"es": "Numeración de miembros físicos continuos", "en": "Numbering of continuous physical members"},
    }
    lang = "en" if idioma in ["ingles", "en"] else "es"



    from docx.enum.table import WD_ROW_HEIGHT

    for idx, num_fig in enumerate([1, 2, 3, 4, 5, 6], start=1):
        if num_fig in image_slots:
            if isinstance(image_slots[num_fig], tuple):
                ruta, _ = image_slots[num_fig]
            else:
                ruta = image_slots[num_fig]

            print(f"✅ Agregando imagen {num_fig}: {ruta}")

            try:
                import os
                if not os.path.exists(ruta):
                    print(f"❌ Archivo no encontrado: {ruta}")
                    continue

                # Crear tabla 1x1 para el borde y ajustar tamaño exacto de la celda
                tabla = nuevo.add_table(rows=1, cols=1)
                tabla.alignment = 1  # Centrar tabla
                tabla.autofit = False  # Desactivar ajuste automático
                tabla.columns[0].width = Cm(17)
                tabla.rows[0].height = Cm(12)
                celda = tabla.cell(0, 0)
                agregar_borde_celda(celda, color="000000", size=4)  # 0.5pt

                # Ajustar tamaño de la celda exactamente al tamaño de la imagen
                celda.width = Cm(17)
                celda.height = Cm(12)
                tabla.columns[0].width = Cm(17)
                tabla.rows[0].height = Cm(12)

                # Insertar imagen en la celda, ajustando el tamaño para que coincida con el borde
                run = celda.paragraphs[0].add_run()
                run.add_picture(ruta, width=Cm(17), height=Cm(12))
                # Ajustar el espaciado y márgenes para que la imagen quede compacta y centrada
                celda.paragraphs[0].paragraph_format.left_indent = 0
                celda.paragraphs[0].paragraph_format.right_indent = 0
                celda.paragraphs[0].paragraph_format.space_after = Pt(0)
                celda.paragraphs[0].paragraph_format.space_before = Pt(0)

                # Espacio después de la tabla
                nuevo.add_paragraph()

                # Título de la figura
                titulo = titulos[num_fig][lang]
                p2 = nuevo.add_paragraph(f"{'Figure' if lang == 'en' else 'Figura'} {idx}. {titulo}")
                p2.alignment = 1
                r2 = p2.runs[0]
                r2.font.name = "Arial"
                r2.font.size = Pt(12)

                # Salto de página después de cada imagen, menos la última
                if idx < 6:
                    nuevo.add_page_break()

            except Exception as e:
                print(f"❌ Error al agregar imagen {num_fig}: {e}")
                continue

    # Al final de la sección, forzar salto de página para que la siguiente sección siempre empiece en nueva hoja
    nuevo.add_page_break()

    print(f"✅ Sección 1 completada con {len([n for n in [1,2,3,4,5,6] if n in image_slots])} imágenes")
    return nuevo

    