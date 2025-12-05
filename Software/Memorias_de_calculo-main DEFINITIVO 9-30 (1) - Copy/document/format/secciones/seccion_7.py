from docx.shared import Pt, Cm
from docx.oxml import parse_xml
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
import os
import sys
import pandas as pd
import re

# Ajusta estas rutas según tu proyecto
project_root = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..', '..'))
sys.path.append(project_root)
sys.path.append(os.path.join(project_root, 'scripts'))
sys.path.append(os.path.join(project_root, 'componentes'))

from scripts.extract_tables_of_excel import extract_tables_from_excel_verificacion_por_deflexion
from componentes.poner_bordes_tablas import poner_bordes_tabla
from componentes.set_repeat_header import set_repeat_table_header
from document.format.helpers import format_number

def merge_empty_rows_across_table(table):
    """Fusiona filas consecutivas completamente vacías en una sola celda que abarca todas las columnas."""
    if len(table.rows) <= 1:
        return
    i = 1  # empieza después del encabezado
    while i < len(table.rows):
        # detectar inicio de secuencia de filas vacías
        if all(cell.text.strip() == "" for cell in table.rows[i].cells):
            start = i
            end = i
            # buscar siguientes vacías
            while end + 1 < len(table.rows) and all(cell.text.strip() == "" for cell in table.rows[end + 1].cells):
                end += 1
            # si hay más de una, fusionar desde start hasta end
            if end > start:
                first = table.cell(start, 0)
                last = table.cell(end, len(table.columns) - 1)
                first.merge(last)
            i = end + 1
        else:
            i += 1

def verificacion_de_deflexion_vertical_en_vigas(nuevo, idioma, file_path):
    """
    Genera la sección de verificación de deflexiones verticales en vigas
    
    Args:
        nuevo: Documento de Word donde agregar la sección
        idioma: Idioma para los textos ('es' o 'en')
        file_path: Ruta del archivo Excel con los datos
    """
    print(f"[DEBUG] Generando sección 7 - Idioma: {idioma}, Archivo: {file_path}")
    
    key = 'en' if idioma.lower() in ('ingles', 'en') else 'es'
    textos = {
        'intro': {
            'en': "With the design spectrum already scaled, the structure is designed to meet the following regulatory requirements:",
            'es': "Con el espectro de diseño ya escalado se diseña la estructura para que cumpla con los siguientes requisitos normativos:"
        },
        'titulo_sec': {
            'en': "Verification of deflections",
            'es': "Verificación de deflexiones"
        },
        'previo_tabla': {
            'en': "The vertical and horizontal deflection in the beams must comply with project criteria:",
            'es': "La deflexión vertical y horizontal en las vigas de la estructura debe ser menor a lo indicado en los criterios de diseño del proyecto:"
        },
        'titulo_tab': {
            'en': "Table to verify vertical deflection in beams",
            'es': "Tabla para verificar la deflexión vertical en vigas"
        },
        'sin_datos': {
            'en': "No vertical deflection verification data found.",
            'es': "No se encontraron datos de verificación de deflexión vertical."
        }
    }

    # Introducción y títulos
    nuevo.add_paragraph()
    p = nuevo.add_paragraph(textos['intro'][key])
    p.runs[0].font.name = "Arial"; p.runs[0].font.size = Pt(12)

    nuevo.add_paragraph()
    p = nuevo.add_paragraph(textos['titulo_sec'][key])
    r = p.runs[0]; r.font.name = "Arial"; r.font.size = Pt(12); r.bold = True
    p.alignment = 0
    nuevo.add_paragraph()

    p = nuevo.add_paragraph(textos['previo_tabla'][key])
    p.runs[0].font.name = "Arial"; p.runs[0].font.size = Pt(12)
    nuevo.add_paragraph()

    p = nuevo.add_paragraph(textos['titulo_tab'][key])
    r = p.runs[0]; r.font.name = "Arial"; r.font.size = Pt(12); r.bold = True
    p.alignment = 0
    nuevo.add_paragraph()

    # Extracción
    try:
        verif_dict = extract_tables_from_excel_verificacion_por_deflexion(file_path)
    except Exception as e:
        print(f"Error extrayendo datos de Excel para sección 7: {e}")
        nuevo.add_paragraph(textos['sin_datos'][key])
        nuevo.add_paragraph()
        nuevo.add_page_break()
        return
        
    # Buscar la primera hoja con datos
    df_verif = None
    for k in verif_dict:
        if verif_dict[k] is not None and not verif_dict[k].empty:
            df_verif = verif_dict[k]
            break
    
    if df_verif is None or df_verif.empty:
        nuevo.add_paragraph(textos['sin_datos'][key])
        nuevo.add_paragraph()
        nuevo.add_page_break()
        return

    # Mapeos y limpieza

    HEADERS_ES = [
        "ID", "Tipo", "Longitud (m)", "Caso de Carga",
        "Max DY (mm)", "Factor", "Deflexión Permisible L/X (mm)",
        "Cumple Norma"
    ]
    HEADERS_EN = [
        "ID", "Type", "Length (m)", "Load Case",
        "Max DY (mm)", "Factor", "Allowable Deflection L/X (mm)",
        "Standard Compliance"
    ]
    MAP_ES = {
        "ID": "ID",
        "Tipo": "Tipo",
        "Physical Member ID": "ID",
        "Analytical Member ID": "Tipo",
        "Longitud de miembro (m)": "Longitud (m)",
        "Caso de carga crítico": "Caso de Carga",
        "Max DY (mm)": "Max DY (mm)",
        "Factor L/X": "Factor",
        "Deflexión Permisible (mm)": "Deflexión Permisible L/X (mm)",
        "Cumple Norma": "Cumple Norma"
    }
    MAP_EN = {
        "ID": "ID",
        "Tipo": "Type",
        "Physical Member ID": "ID",
        "Analytical Member ID": "Type",
        "Longitud de miembro (m)": "Length (m)",
        "Member Length (m)": "Length (m)",
        "Caso de carga crítico": "Load Case",
        "Critical Load Case": "Load Case",
        "Max DY (mm)": "Max DY (mm)",
        "Factor L/X": "Factor",
        "L/X Factor": "Factor",
        "Deflexión Permisible (mm)": "Allowable Deflection L/X (mm)",
        "Allowable Deflection (mm)": "Allowable Deflection L/X (mm)",
        "Cumple Norma": "Standard Compliance",
        "Standard Compliance": "Standard Compliance"
    }

    if key == 'en':
        rename_map = {orig: MAP_EN[orig] for orig in df_verif.columns if orig in MAP_EN}
        headers = HEADERS_EN
    else:
        rename_map = {orig: MAP_ES[orig] for orig in df_verif.columns if orig in MAP_ES}
        headers = HEADERS_ES

    # Aplicar renombrado
    df = df_verif.rename(columns=rename_map)

    # Reemplazar valores de 'Miembro Físico' y 'Miembro Analítico'
    if key == 'es':
        if 'Tipo' in df.columns:
            df['Tipo'] = df['Tipo'].replace({'PM': 'MF', 'AM': 'EA'})
    else: # en
        if 'Type' in df.columns:
            df['Type'] = df['Type'].replace({'AM': 'AE'})
    
    # Verificar que tenemos las columnas necesarias
    available_headers = [h for h in headers if h in df.columns]
    if len(available_headers) < 3:  # Mínimo 3 columnas para que tenga sentido
        print(f"Advertencia: Solo se encontraron {len(available_headers)} columnas válidas en sección 7")
        print(f"Columnas disponibles: {list(df.columns)}")
        print(f"Columnas esperadas: {headers}")
    
    # Seleccionar solo las columnas disponibles
    df = df[available_headers]
    df = df.dropna(how='all').copy()
    
    # Agregar columnas faltantes como vacías
    for h in headers:
        if h not in df.columns:
            df[h] = ""
    
    # Reordenar según el orden esperado
    df = df[headers]
    df = df[~df.apply(lambda row: all((str(x).strip()=="" or pd.isna(x)) for x in row), axis=1)]

    # Verificar que tenemos datos después de la limpieza
    if df.empty:
        nuevo.add_paragraph(textos['sin_datos'][key])
        nuevo.add_paragraph()
        nuevo.add_page_break()
        return

    # Anchos fijos para 8 columnas
    n = len(headers)
    if n == 8:
        col_widths = [2.0, 2.21, 2.11, 1.5, 1.75, 1.61, 2.5, 2.0]
    else:
        usable = 19.57
        col_widths = [usable / n] * n

    # Crear tabla
    tabla = nuevo.add_table(rows=1, cols=n)
    set_repeat_table_header(tabla.rows[0])
    tabla.autofit = False
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabla.style = None

    # Encabezados con letra tamaño 9 igual que sección 8 (sin saltos de línea)
    for j, h in enumerate(headers):
        cell = tabla.cell(0, j)
        cell.text = h
        p0 = cell.paragraphs[0]
        p0.alignment = 1
        p0.paragraph_format.left_indent = 0
        p0.paragraph_format.right_indent = 0
        p0.paragraph_format.space_after = 0
        p0.paragraph_format.space_before = 0
        while len(p0.runs) > 1:
            p0.runs[-1].clear(); del p0.runs[-1]
        run = p0.runs[0]
        run.text = h; run.font.name = "Arial"; run.font.size = Pt(9); run.font.bold = True
        rPr = run._element.get_or_add_rPr()
        sz = parse_xml(r'<w:sz xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="18"/>')
        szCs = parse_xml(r'<w:szCs xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="18"/>')
        rPr.append(sz); rPr.append(szCs)
        shade = parse_xml(r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="D9D9D9"/>')
        cell._tc.get_or_add_tcPr().append(shade)
        cell.width = Cm(col_widths[j]); cell.vertical_alignment = 1

    tabla.rows[0].height = Cm(1.6)
    tabla.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    # Filas de datos
    for row in df.itertuples(index=False):
        cells = tabla.add_row().cells
        for j, val in enumerate(row):
            text = ""
            if pd.notna(val):
                if isinstance(val, (int, float)):
                    text = format_number(val)
                else:
                    text = str(val).strip()
            # Traducción de "si"/"no" a "Yes"/"No" en inglés
            if key == 'en' and headers[j] == "Standard Compliance":
                if text.lower() in ["si", "sí", "yes", "true", "1"]:
                    text = "Yes"
                elif text.lower() in ["no", "false", "0"]:
                    text = "No"
            if key == 'es' and headers[j] == "Cumple Norma":
                if text.lower() in ["yes", "true", "1"]:
                    text = "Sí"
                elif text.lower() in ["no", "false", "0"]:
                    text = "No"
            cell = cells[j]; cell.text = text
            p1 = cell.paragraphs[0]
            p1.alignment = 1
            p1.paragraph_format.left_indent = 0
            p1.paragraph_format.right_indent = 0
            p1.paragraph_format.space_after = 0
            p1.paragraph_format.space_before = 0
            while len(p1.runs) > 1:
                p1.runs[-1].clear(); del p1.runs[-1]
            run1 = p1.runs[0]
            run1.text = text; run1.font.name = "Arial"; run1.font.size = Pt(10)
            # Sombreado verde si cumple norma, rojo si NO
            if headers[j] in ["Cumple Norma", "Standard Compliance"]:
                if text.lower() in ["si","sí","yes","true","1"]:
                    shade = parse_xml(r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="C6EFCE"/>')
                    cell._tc.get_or_add_tcPr().append(shade)
                elif text.lower() in ["no","false","0"]:
                    shade = parse_xml(r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="FFC7CE"/>')
                    cell._tc.get_or_add_tcPr().append(shade)
            cell.width = Cm(col_widths[j]); cell.vertical_alignment = 1
        tabla.rows[-1].height = Cm(0.58)
        tabla.rows[-1].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    # Bordes
    poner_bordes_tabla(tabla)

    # Fusionar filas vacías consecutivas
    merge_empty_rows_across_table(tabla)

    nuevo.add_paragraph()
    nuevo.add_page_break()