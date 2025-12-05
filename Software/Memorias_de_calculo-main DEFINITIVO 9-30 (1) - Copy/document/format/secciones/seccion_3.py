from docx.shared import Pt, Cm
from docx.oxml import parse_xml
from docx.enum.table import WD_ROW_HEIGHT_RULE
import os
import sys
project_root = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..', '..'))
sys.path.append(project_root)
sys.path.append(os.path.join(project_root, 'staad_automation'))
sys.path.append(os.path.join(project_root, 'componentes'))
from componentes.poner_bordes_tablas import poner_bordes_tabla
from componentes.set_repeat_header import set_repeat_table_header
from document.format.helpers import format_number

def casos_de_carga_primarias(nuevo, idioma):
    # Diccionario de traducción para la columna Tipo
    traduccion_tipo = {
        "Dead": "Muerta",
        "Viva": "Viva",
        "Viva de techo": "Viva de techo",
        "Temperatura": "Temperatura",
        "Wind": "Viento",
        "Seismic-H": "Sismo-H",
        "Seismic-V": "Sismo-V"
    }
    traduccion_tipo_en = {
        "Muerta": "Dead",
        "Viva": "Live",
        "Viva de techo": "Roof Live",
        "Temperatura": "Temperature",
        "Viento": "Wind",
        "Sismo-H": "Seismic-H",
        "Sismo-V": "Seismic-V"
    }

    # Diccionario de traducción para la columna Descripción
    traduccion_desc = {
        "Peso propio (D)": "Self-weight (D)",
        "Peso en Operación (OP)": "Operating weight (OP)",
        "Carga Prueba hidrostática (F)": "Hydrostatic test load (F)",
        "Carga viva (L)": "Live load (L)",
        "Carga viva de techo (Lr)": "Roof live load (Lr)",
        "Cargas por temperatura (TST)": "Temperature loads (TST)",
        "Carga por fricción de tuberías (TT)": "Pipe friction loads (TT)",
        "Fuerzas por anclaje de tuberías (TS)": "Pipe anchorage forces (TS)",
        "Viento en dirección X (WX)": "Wind in X direction (WX)",
        "Viento en dirección -X (WX)": "Wind in -X direction (WX)",
        "Viento en dirección Z (WZ)": "Wind in Z direction (WZ)",
        "Viento en dirección -z (WZ)": "Wind in -Z direction (WZ)",
        "Sismo en dirección X (EX)": "Seismic in X direction (EX)",
        "Sismo en dirección Z (EZ)": "Seismic in Z direction (EZ)",
        "Sismo en dirección Y (EV)": "Seismic in Y direction (EV)"
    }
    traduccion_desc_es = {
        "Self-weight (D)": "Peso propio (D)",
        "Operating weight (OP)": "Peso en Operación (OP)",
        "Hydrostatic test load (F)": "Carga Prueba hidrostática (F)",
        "Live load (L)": "Carga viva (L)",
        "Roof live load (Lr)": "Carga viva de techo (Lr)",
        "Temperature loads (TST)": "Cargas por temperatura (TST)",
        "Pipe friction loads (TT)": "Carga por fricción de tuberías (TT)",
        "Pipe anchorage forces (TS)": "Fuerzas por anclaje de tuberías (TS)",
        "Wind in X direction (WX)": "Viento en dirección X (WX)",
        "Wind in -X direction (WX)": "Viento en dirección -X (WX)",
        "Wind in Z direction (WZ)": "Viento en dirección Z (WZ)",
        "Wind in -Z direction (WZ)": "Viento en dirección -z (WZ)",
        "Seismic in X direction (EX)": "Sismo en dirección X (EX)",
        "Seismic in Z direction (EZ)": "Sismo en dirección Z (EZ)",
        "Seismic in Y direction (EV)": "Sismo en dirección Y (EV)"
    }
    nuevo.add_page_break()

    parrafo_carga = nuevo.add_paragraph()
    if idioma in ["ingles", "en"]:
        parrafo_carga.add_run("According to the project design criteria, the primary loads used in the analysis are detailed below").font.name = "Arial"
        parrafo_carga.add_run(" They are not included in this report.").font.name = "Arial"
    else:
        parrafo_carga.add_run("Conforme a los criterios de diseño del proyecto, las cargas primarias utilizadas en el análisis se detallan a continuación:").font.name = "Arial"
        parrafo_carga.add_run(" No están incluidas en este informe.").font.name = "Arial"        

    nuevo.add_paragraph()    

    # Título según idioma
    if idioma in ["ingles", "en"]:
        nuevo.add_paragraph("Primary load table")
        p = nuevo.paragraphs[-1]
        run = p.runs[0] if p.runs else p.add_run() 
        run.font.name = "Arial"
        run.font.size = Pt(12)
        run.bold = True
        p.alignment = 0
        nuevo.add_paragraph()
    else:
        nuevo.add_paragraph("Tabla de cargas primarias")
        p = nuevo.paragraphs[-1]
        run = p.runs[0] if p.runs else p.add_run()
        run.font.name = "Arial"
        run.font.size = Pt(12)
        run.bold = True 
        p.alignment = 0

    nuevo.add_paragraph()

    # Intentar extraer tabla de combinaciones de carga
    try:
        from staad_automation.extract_load_primary import extract_primary_loads  # <--- AQUÍ DENTRO
        primary_loads = extract_primary_loads()
    except Exception as e:
        print(f"⚠️ No se pudo extraer cargas primarias: {e}")
        primary_loads = None

    if primary_loads:
        # Diccionarios de traducción para los valores más comunes
        tipo_trad = {
            # Español a inglés
            "dead": "DEAD", "peso propio (d)": "DEAD", "peso t.operacion (op)": "DEAD", "peso carga prueba hidrostatica (h)": "DEAD",
            "live": "LIVE", "carga viva": "LIVE",
            "carga": "LOAD", "termica estructura (tst)": "THERMAL STRUCTURE (TST)",
            "fuerzas": "FORCES", "de friccion tuberias (tt)": "PIPE FRICTION FORCES (TT)", "de anclaje tuberias (ts)": "PIPE ANCHOR FORCES (TS)",
            "viento": "WIND", "direccion x": "DIRECTION X", "direccion -x": "DIRECTION -X", "direccion z": "DIRECTION Z", "direccion -z": "DIRECTION -Z",
            "sismo": "SEISMIC", "dir x": "DIR X", "dir z": "DIR Z", "dir y (vertical)": "DIR Y (VERTICAL)"
        }
        tipo_trad_inv = {v.lower(): k.capitalize() for k, v in tipo_trad.items()}  # Inglés a español

        def traducir(texto, idioma_destino):
            t = texto.strip().lower()
            if idioma_destino == 'en':
                return tipo_trad.get(t, texto)
            else:
                return tipo_trad_inv.get(t, texto)

        # Crear tabla con encabezados según idioma
        if idioma in ["ingles", "en"]:
            headers = ["Case", "Type", "Description"]
        else:
            headers = ["Caso", "Tipo", "Descripción"]

        ncols = len(headers)
        ancho_total = 16.5  # Ajusta según el margen de tu documento Word
        col_widths = [3.0, 3.0, ancho_total - 6.0]  # Caso, Tipo, Descripción
        tabla_cargas = nuevo.add_table(rows=1, cols=ncols)
        set_repeat_table_header(tabla_cargas.rows[0])
        tabla_cargas.autofit = False
        tabla_cargas.alignment = 1  # Centrar la tabla
        for i, h in enumerate(headers):
            cell = tabla_cargas.cell(0, i)
            cell.text = h
            p = cell.paragraphs[0]
            p.alignment = 1
            p.runs[0].font.bold = True
            p.runs[0].font.size = Pt(12)
            shading_elm = parse_xml(
                r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="D9D9D9"/>'
            )
            cell._tc.get_or_add_tcPr().append(shading_elm)
            cell.width = Cm(col_widths[i])

        row_height = Cm(0.8)
        for caso, tipo, desc in primary_loads:
            row = tabla_cargas.add_row()
            row.height = row_height
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            row.cells[0].text = format_number(caso)
            row.cells[1].text = tipo
            row.cells[2].text = str(desc)
            for i in range(ncols):
                p = row.cells[i].paragraphs[0]
                p.alignment = 1
                p.runs[0].font.size = Pt(12)
                row.cells[i].width = Cm(col_widths[i])
        poner_bordes_tabla(tabla_cargas)
    else:
        nuevo.add_paragraph()
        if idioma in ["ingles", "en"]:
            nuevo.add_paragraph("Primary load cases not available.")
        else:
            nuevo.add_paragraph("Casos de carga no disponibles.")
    nuevo.add_paragraph()
    
    # Salto de página al final de la sección
    nuevo.add_page_break()