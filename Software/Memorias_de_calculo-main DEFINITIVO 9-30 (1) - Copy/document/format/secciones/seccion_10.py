import re
from docx.shared import Pt, Cm
from docx.oxml import parse_xml
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
import os
import sys
import openpyxl

project_root = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..', '..'))
sys.path.append(project_root)
sys.path.append(os.path.join(project_root, 'scripts'))
sys.path.append(os.path.join(project_root, 'componentes'))

from componentes.poner_bordes_tablas import poner_bordes_tabla
from componentes.set_repeat_header import set_repeat_table_header
from document.format.helpers import format_number, format_decimal_2
from document.format.secciones.table_creator import create_table_with_merged_cells

def find_and_create_tables(nuevo, sheet, table_title_keywords, key):
    tables_found = 0
    for row_idx, row in enumerate(sheet.iter_rows(), 1):
        for cell in row:
            if cell.value and isinstance(cell.value, str):
                for keyword in table_title_keywords:
                    if keyword.lower() in cell.value.lower():
                        tables_found += 1
                        print(f"Found table '{cell.value}' at row {row_idx}")

                        # Add table title to the document
                        title = cell.value
                        if key == 'en':
                            title = re.sub(r'verificaci(ó|o)n de derivas', 'Drift Check', title, flags=re.IGNORECASE)
                            title = re.sub(r'por sismo', 'by Earthquake', title, flags=re.IGNORECASE)
                            title = re.sub(r'dirección', 'Direction', title, flags=re.IGNORECASE)
                            title = re.sub(r'combinaciones de servicio', 'Service Combinations', title, flags=re.IGNORECASE)
                            title = re.sub(r'combinaciones de (ú|u)ltimas', 'Ultimate Combinations', title, flags=re.IGNORECASE)
                        p = nuevo.add_paragraph(title)
                        r = p.runs[0]; r.font.name = "Arial"; r.font.size = Pt(12); r.bold = True
                        p.alignment = 0
                        nuevo.add_paragraph()

                        # Find header row (first non-empty row after title)
                        header_row_idx = row_idx + 1
                        while header_row_idx <= sheet.max_row and not any(sheet.cell(row=header_row_idx, column=j).value for j in range(1, sheet.max_column + 1)):
                            header_row_idx += 1

                        if header_row_idx > sheet.max_row:
                            continue

                        # Read headers from the sheet
                        headers = []
                        start_col = -1
                        for j in range(1, sheet.max_column + 1):
                            header_cell = sheet.cell(row=header_row_idx, column=j)
                            if header_cell.value:
                                if start_col == -1:
                                    start_col = j
                                headers.append(str(header_cell.value))
                            elif start_col != -1:
                                break
                        
                        num_cols = len(headers)
                        if num_cols == 0:
                            continue

                        # Find and read data rows
                        start_row_data = header_row_idx + 1
                        num_data_rows = 0
                        for i in range(start_row_data, sheet.max_row + 1):
                            row_values = [sheet.cell(row=i, column=j).value for j in range(start_col, start_col + num_cols)]
                            if any(v is not None for v in row_values):
                                num_data_rows += 1
                            elif all(v is None for v in row_values):
                                break # Stop at first fully empty row
                        
                        # Create and populate table
                        tabla = nuevo.add_table(rows=num_data_rows + 1, cols=num_cols)
                        set_repeat_table_header(tabla.rows[0])
                        tabla.autofit = False
                        tabla.alignment = WD_TABLE_ALIGNMENT.CENTER

                        create_table_with_merged_cells(sheet, tabla, header_row_idx, num_data_rows + 1, start_col, num_cols, key, headers)

                        poner_bordes_tabla(tabla)
                        nuevo.add_paragraph()
                        break 
    return tables_found > 0

def verificacion_por_sismo(nuevo,idioma,file_path):
    key = 'en' if idioma.lower() in ('ingles', 'en') else 'es'
    textos = {
        'titulo': {
            'en': "Verification of earthquake displacements",
            'es': "Verificación de desplazamientos por sismo"
        },
        'intro': {
            'en': "The lateral drift ratio in the structure's columns due to seismic loads must be less than that indicated in the standard governing the design and that indicated in the project's design criteria.",
            'es': "La relación de deriva lateral en las columnas de la estructura debido a las cargas sísmicas debe ser menor a lo indicado en la norma que rige el diseño y la indicada en los criterios de diseño del proyecto."
        },
        'sin_datos': {
            'en': "No earthquake displacement verification data found.",
            'es': "No se encontraron datos de verificación de desplazamientos por sismo."
        }
    }

    nuevo.add_paragraph()
    p = nuevo.add_paragraph(textos['titulo'][key])
    r = p.runs[0]; r.font.name = "Arial"; r.font.size = Pt(12); r.bold = True
    p.alignment = 0
    nuevo.add_paragraph()
    p = nuevo.add_paragraph(textos['intro'][key])
    p.runs[0].font.name = "Arial"; p.runs[0].font.size = Pt(12)
    nuevo.add_paragraph()

    try:
        workbook = openpyxl.load_workbook(file_path, data_only=True)
        sheet = workbook['Sismo']
    except (FileNotFoundError, KeyError) as e:
        print(f"Error opening Excel or sheet 'Sismo': {e}")
        nuevo.add_paragraph(textos['sin_datos'][key])
        return

    table_title_keywords = ["Verificación de derivas", "Verificacion de derivas"]
    if not find_and_create_tables(nuevo, sheet, table_title_keywords, key):
        nuevo.add_paragraph(textos['sin_datos'][key])

    nuevo.add_page_break()