from docx.shared import Pt
import os
import sys
project_root = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..', '..'))
sys.path.append(project_root)
sys.path.append(os.path.join(project_root, 'scripts'))
sys.path.append(os.path.join(project_root, 'componentes'))
sys.path.append(os.path.join(project_root, 'staad_automation'))
from staad_automation.get_path_of_staad_connetc import get_path_of_staad_connect, transforma_el_std_a_txt

def leer_std_a_texto(nuevo, idioma):
    import pythoncom
    pythoncom.CoInitialize()
    nuevo.add_page_break()
    # Título según idioma
    if idioma in ["ingles", "en"]:
        nuevo.add_paragraph("Data Entry")
        p = nuevo.paragraphs[-1]
        run = p.runs[0] if p.runs else p.add_run()
        run.font.name = "Arial"
        run.font.size = Pt(12)
        run.bold = True
        p.alignment = 0
        nuevo.add_paragraph()
    else:
        nuevo.add_paragraph("Entrada de Datos")
        p = nuevo.paragraphs[-1]
        run = p.runs[0] if p.runs else p.add_run()
        run.font.name = "Arial"
        run.font.size = Pt(12)
        run.bold = True
        p.alignment = 0
        nuevo.add_paragraph()
        
    # --- Sección para insertar el contenido del archivo .std como texto ---
    try:
        path_staad = get_path_of_staad_connect()
        if path_staad and os.path.isfile(path_staad):
            # Transforma y obtiene la ruta del archivo .txt
            path_txt = transforma_el_std_a_txt(path_staad)
            # Lee el contenido del archivo .std (no del .txt)
            with open(path_staad, 'r') as file:
                contenido_std = file.read()
            # Insertar el contenido como un párrafo preformateado
            nuevo.add_paragraph()  # Espacio antes
            run = p.runs[0] if p.runs else p.add_run()
            run.font.name = "Arial"
            run.font.size = Pt(12)
            run.bold = True
            p.alignment = 0
            nuevo.add_paragraph()  # Espacio

            # Insertar el contenido del archivo como texto preformateado
            for linea in contenido_std.splitlines():
                p_txt = nuevo.add_paragraph(linea)
                p_txt.style = 'Normal'
                for run in p_txt.runs:
                    run.font.name = "Consolas"
                    run.font.size = Pt(8)
            nuevo.add_paragraph()  # Espacio después
        else:
            nuevo.add_paragraph("Proyecto STAAD no encontrado.")
    except Exception:
        nuevo.add_paragraph("Proyecto STAAD no encontrado.")