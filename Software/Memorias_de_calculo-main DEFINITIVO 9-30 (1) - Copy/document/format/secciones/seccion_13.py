from docx.shared import Pt
from docx.oxml import parse_xml
import pandas as pd
import os
import sys
project_root = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..', '..'))
sys.path.append(project_root)
sys.path.append(os.path.join(project_root, 'scripts'))
sys.path.append(os.path.join(project_root, 'componentes'))

from scripts.extract_tables_of_excel import extract_tables_from_excel_computos
from componentes.poner_bordes_tablas import poner_bordes_tabla
from componentes.set_repeat_header import set_repeat_table_header
from document.format.helpers import format_number



def computos_metricos(nuevo, idioma, file_path):
        nuevo.add_paragraph()
        # Título según idioma
        if idioma in ["ingles", "en"]:
            nuevo.add_paragraph("Structure Bill of Quantities")
            p = nuevo.paragraphs[-1]
            run = p.runs[0] if p.runs else p.add_run()
            run.font.name = "Arial"
            run.font.size = Pt(12)
            run.bold = True
            p.alignment = 0
            nuevo.add_paragraph()
        else:
            nuevo.add_paragraph("Cómputos métricos de la estructura")
            p = nuevo.paragraphs[-1]
            run = p.runs[0] if p.runs else p.add_run()
            run.font.name = "Arial"
            run.font.size = Pt(12)
            run.bold = True
            p.alignment = 0
            nuevo.add_paragraph()

        # Extraer tabla de cómputos métricos
        try:
            computos_dict = extract_tables_from_excel_computos(file_path)
            df_computos = computos_dict.get('CP', None)
        except Exception as e:
            df_computos = None

        # Encabezados y mapeo
        HEADERS_ES = ["Tipo", "Perfil", "Longitud (m)", "Peso (kN)"]
        HEADERS_EN = ["Type", "Profile", "Length (m)", "Weight (kN)"]
        COL_MAP = {
            "Tipo": ["Tipo de Perfil", "Tipo"],
            "Type": ["Type", "Tipo de Perfil", "Tipo"],
            "Perfil": ["Perfil", "Profile"],
            "Profile": ["Profile", "Perfil"],
            "Longitud (m)": ["Longitud Total (m)", "Longitud (m)", "Length (m)"],
            "Length (m)": ["Length (m)", "Longitud Total (m)", "Longitud (m)"],
            "Peso (kN)": ["Peso Total (kN)", "Peso (kN)", "Weight (kN)"],
            "Weight (kN)": ["Weight (kN)", "Peso Total (kN)", "Peso (kN)"]
        }

        if df_computos is not None and not df_computos.empty:
            headers = HEADERS_EN if idioma in ["ingles", "en"] else HEADERS_ES
            cols = []
            for h in headers:
                found = None
                for c in COL_MAP[h]:
                    if c in df_computos.columns:
                        found = c
                        break
                if found:
                    cols.append(found)
            if cols:
                df_show = df_computos[cols].copy()
                df_show.columns = headers[:len(cols)]
                ncols = len(df_show.columns)
                nrows = len(df_show)
                tabla_computos = nuevo.add_table(rows=1, cols=ncols)
                set_repeat_table_header(tabla_computos.rows[0])
                tabla_computos.autofit = False
                from docx.shared import Cm
                from docx.enum.table import WD_TABLE_ALIGNMENT
                tabla_computos.alignment = WD_TABLE_ALIGNMENT.CENTER
                col_widths = [2.93, 2.01, 3.99, 3.35]
                # Encabezados
                for j, col in enumerate(df_show.columns):
                    cell = tabla_computos.cell(0, j)
                    cell.text = str(col)
                    p = cell.paragraphs[0]
                    p.alignment = 1
                    run = p.runs[0] if p.runs else p.add_run()
                    run.font.bold = True
                    run.font.size = Pt(11)
                    run.font.name = "Arial"
                    shading_elm = parse_xml(
                        r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="D9D9D9"/>'
                    )
                    cell._tc.get_or_add_tcPr().append(shading_elm)
                    if j < len(col_widths):
                        cell.width = Cm(col_widths[j])
                # Filas de datos
                for i, (_, row) in enumerate(df_show.iterrows()):
                    row_cells = tabla_computos.add_row().cells
                    for j, val in enumerate(row):
                        cell = row_cells[j]
                        if pd.isna(val):
                            cell.text = ""
                        else:
                            if isinstance(val, (int, float)):
                                cell.text = format_number(val)
                            else:
                                cell.text = str(val).strip()
                        p = cell.paragraphs[0]
                        p.alignment = 1
                        run = p.runs[0] if p.runs else p.add_run()
                        run.font.size = Pt(10)
                        run.font.name = "Arial"
                        if j < len(col_widths):
                            cell.width = Cm(col_widths[j])
                poner_bordes_tabla(tabla_computos)
                nuevo.add_paragraph()
        else:
            if idioma in ["ingles", "en"]:
                nuevo.add_paragraph("No bill of quantities data found.")
            else:
                nuevo.add_paragraph("No se encontraron datos de cómputos métricos.")
            nuevo.add_paragraph()