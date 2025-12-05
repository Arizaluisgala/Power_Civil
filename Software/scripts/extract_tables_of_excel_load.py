#!/usr/bin/env python3
import pandas as pd
import os
import tkinter as tk
from tkinter import filedialog
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from docx.oxml import parse_xml

# Patrones exactos que deben coincidir
TARGET_PATTERNS = [
    ['L/C', 'Direction', 'Load', 'Min X', 'Max X', 'Min Y', 'Max Y', 'Min Z', 'Max Z'],
    ['L/C', 'Parameter', 'Value'],
    ['L/C', 'Beam', 'Type', 'Axial', 'Cross', 'Side', 'Elongation', 'Strain Rate'],
    ['L/C', 'Beam', 'Type', 'Direction', 'Fa', 'Da', 'Fb', 'Db', 'Ecc.'],
    ['L/C', 'Node', 'FX', 'FY', 'FZ', 'MX', 'MY', 'MZ', 'Type'],
    ['L/C', 'Direction', 'Factor', 'Assigned Geometry'],
    ['Type', 'Title', 'Intensity', 'Height']
]

# Diccionario de títulos por patrón
TITULOS_TABLAS = {     0: {"es": "Definiciones de carga de viento", "en": "Wind Load Definitions"},     1: {"es": "Pesos propios", "en": "Selfweights"},     2: {"es": "Cargas nodales", "en": "Nodal Loads"},     3: {"es": "Cargas de viga", "en": "Beam Loads"},     4: {"es": "Cargas de temperatura", "en": "Temperature Loads"},     5: {"es": "Parámetros del espectro", "en": "Spectrum Parameters"},     6: {"es": "Cargas de piso unidireccionales", "en": "One-Way Floor Loads"},  }
# Diccionario de traducción de encabezados
TRADUCCION_HEADERS = {
    "L/C": {"es": "Caso de Carga", "en": "L/C"},
    "Direction": {"es": "Dirección", "en": "Direction"},
    "Load": {"es": "Carga", "en": "Load"},
    "Min X": {"es": "Mín X", "en": "Min X"},
    "Max X": {"es": "Máx X", "en": "Max X"},
    "Min Y": {"es": "Mín Y", "en": "Min Y"},
    "Max Y": {"es": "Máx Y", "en": "Max Y"},
    "Min Z": {"es": "Mín Z", "en": "Min Z"},
    "Max Z": {"es": "Máx Z", "en": "Max Z"},
    "Parameter": {"es": "Parámetro", "en": "Parameter"},
    "Value": {"es": "Valor", "en": "Value"},
    "Beam": {"es": "Viga", "en": "Beam"},
    "Type": {"es": "Tipo", "en": "Type"},
    "Axial": {"es": "Axial", "en": "Axial"},
    "Cross": {"es": "Transversal", "en": "Cross"},
    "Side": {"es": "Lado", "en": "Side"},
    "Elongation": {"es": "Elongación", "en": "Elongation"},
    "Strain Rate": {"es": "Tasa de Deformación", "en": "Strain Rate"},
    "Fa": {"es": "Fa", "en": "Fa"},
    "Da": {"es": "Da", "en": "Da"},
    "Fb": {"es": "Fb", "en": "Fb"},
    "Db": {"es": "Db", "en": "Db"},
    "Ecc.": {"es": "Excentricidad", "en": "Ecc."},
    "Node": {"es": "Nodo", "en": "Node"},
    "FX": {"es": "FX", "en": "FX"},
    "FY": {"es": "FY", "en": "FY"},
    "FZ": {"es": "FZ", "en": "FZ"},
    "MX": {"es": "MX", "en": "MX"},
    "MY": {"es": "MY", "en": "MY"},
    "MZ": {"es": "MZ", "en": "MZ"},
    "Factor": {"es": "Factor", "en": "Factor"},
    "Assigned Geometry": {"es": "Geometría Asignada", "en": "Assigned Geometry"},
    "Title": {"es": "Título", "en": "Title"},
    "Intensity": {"es": "Intensidad", "en": "Intensity"},
    "Height": {"es": "Altura", "en": "Height"}
}

def clean_header(row):
    """Limpia una fila de encabezados"""
    return [str(cell).strip() for cell in row if pd.notna(cell) and str(cell).strip() != '']

def select_excel_file():
    """Abre un diálogo para que el usuario seleccione un archivo Excel."""
    root = tk.Tk()
    root.withdraw()
    file_path = filedialog.askopenfilename(
        title="Selecciona el archivo Excel",
        initialdir=os.path.expanduser("~/Desktop"),
        filetypes=[("Archivos Excel", "*.xlsx *.xls"), ("Todos los archivos", "*.*")]
    )
    root.destroy()
    return file_path if file_path else None

def select_sheet_name(excel_file):
    """Permite seleccionar la hoja si hay varias."""
    try:
        xls = pd.ExcelFile(excel_file)
        sheets = xls.sheet_names
        if len(sheets) == 1:
            return sheets[0]
        print("\nHojas disponibles:")
        for i, name in enumerate(sheets, 1):
            print(f"{i}. {name}")
        choice = input("Selecciona el número de hoja (ENTER=1): ").strip()
        return sheets[int(choice) - 1] if choice.isdigit() and 0 < int(choice) <= len(sheets) else sheets[0]
    except Exception as e:
        print("❌ Error leyendo hojas:", e)
        return None

def extract_target_tables(excel_path, sheet_name=0):
    """Extrae las tablas que coinciden con los patrones definidos"""
    print("📖 Leyendo Excel...")
    df = pd.read_excel(excel_path, sheet_name=sheet_name, header=None)
    tables = []
    
    print("🔍 Buscando patrones de tabla...")
    
    for i, row in df.iterrows():
        cleaned_row = clean_header(row)
        
        for pattern_idx, pattern in enumerate(TARGET_PATTERNS):
            if cleaned_row == pattern:
                print(f"   ✅ Patrón {pattern_idx + 1} encontrado en fila {i}")
                
                # Obtener los índices de columnas que tienen datos
                col_indices = []
                for j, val in enumerate(row):
                    if pd.notna(val) and str(val).strip() != '':
                        col_indices.append(j)
                
                # Extraer datos de la tabla con lógica mejorada
                data_rows = []
                current_row = i + 1
                consecutive_empty_rows = 0
                max_empty_rows = 3  # Máximo de filas vacías consecutivas antes de parar
                
                # Buscar datos hasta 50 filas después del encabezado como máximo
                max_search_rows = min(len(df), i + 50)
                
                while current_row < max_search_rows:
                    # Obtener datos de las columnas relevantes
                    data_row = df.iloc[current_row, col_indices].values
                    
                    # Verificar si la fila tiene datos útiles
                    has_data = any(pd.notna(val) and str(val).strip() != '' for val in data_row)
                    
                    if has_data:
                        # Reset contador de filas vacías y agregar datos
                        consecutive_empty_rows = 0
                        data_rows.append(data_row)
                    else:
                        # Contar filas vacías consecutivas
                        consecutive_empty_rows += 1
                        
                        # Si ya tenemos datos y encontramos muchas filas vacías, parar
                        if len(data_rows) > 0 and consecutive_empty_rows >= max_empty_rows:
                            break
                    
                    current_row += 1
                
                # Verificar si encontramos al menos algunos datos
                if not data_rows:
                    # Intentar búsqueda más amplia si no encontramos nada
                    print(f"      🔍 Búsqueda ampliada para patrón en fila {i}...")
                    current_row = i + 1
                    max_search_rows = min(len(df), i + 100)  # Buscar más lejos
                    
                    while current_row < max_search_rows and len(data_rows) < 5:
                        # Buscar cualquier fila que tenga datos en las columnas del patrón
                        full_row = df.iloc[current_row].values
                        
                        # Verificar si hay datos en posiciones similares al patrón
                        potential_data = []
                        for col_idx in col_indices:
                            if col_idx < len(full_row):
                                potential_data.append(full_row[col_idx])
                            else:
                                potential_data.append('')
                        
                        # Si encontramos datos potenciales, incluirlos
                        if any(pd.notna(val) and str(val).strip() != '' for val in potential_data):
                            data_rows.append(potential_data)
                        
                        current_row += 1
                
                # Crear DataFrame con los datos encontrados
                if data_rows:
                    table_df = pd.DataFrame(data_rows, columns=pattern)
                    # Limpiar el DataFrame
                    table_df = table_df.replace('', pd.NA).dropna(how='all')
                    
                    if len(table_df) > 0:  # Solo agregar si tiene datos reales
                        tables.append({
                            "pattern": pattern,
                            "pattern_index": pattern_idx,
                            "start_row": i,
                            "end_row": current_row - 1,
                            "data": table_df
                        })
                        print(f"      📊 {len(table_df)} filas de datos encontradas")
                    else:
                        print(f"      ⚠️  Patrón encontrado pero sin datos válidos")
                else:
                    print(f"      ⚠️  Patrón encontrado pero tabla vacía")
                break
    
    return tables

def poner_bordes_tabla(tabla):
    """Añade bordes a todas las celdas de la tabla"""
    for cell in tabla._cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tcBorders.append(border)
        tcPr.append(tcBorders)

def crear_documento_word_con_tablas(tablas_extraidas, idioma="es"):
    """Crea un documento Word con las tablas extraídas aplicando el estilo especificado"""
    
    print("📝 Creando documento Word...")
    
    # Crear nuevo documento
    doc = Document()
    
    # Título principal
    if idioma.startswith("es"):
        titulo_principal = "Tablas de Cargas Aplicadas"
        texto_intro = "A continuación se presentan las tablas extraídas del análisis estructural:"
    else:
        titulo_principal = "Applied Loads Tables"
        texto_intro = "Below are the tables extracted from the structural analysis:"
    
    # Añadir título principal
    titulo = doc.add_paragraph(titulo_principal)
    titulo_run = titulo.runs[0] if titulo.runs else titulo.add_run()
    titulo_run.font.name = "Arial"
    titulo_run.font.size = Pt(16)
    titulo_run.bold = True
    titulo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Añadir párrafo introductorio
    intro = doc.add_paragraph(texto_intro)
    intro_run = intro.runs[0] if intro.runs else intro.add_run()
    intro_run.font.name = "Arial"
    intro_run.font.size = Pt(12)
    intro.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    doc.add_paragraph()  # Espacio
    
    # Procesar cada tabla
    for tabla in tablas_extraidas:
        # Obtener título de la tabla
        pattern_idx = tabla["pattern_index"]
        titulo_tabla = TITULOS_TABLAS[pattern_idx]["es"] if idioma.startswith("es") else TITULOS_TABLAS[pattern_idx]["en"]
        
        print(f"   📊 Procesando: {titulo_tabla}")
        
        # Añadir título de la tabla
        p_titulo = doc.add_paragraph(titulo_tabla)
        run_titulo = p_titulo.runs[0] if p_titulo.runs else p_titulo.add_run()
        run_titulo.font.name = "Arial"
        run_titulo.font.size = Pt(12)
        run_titulo.bold = True
        p_titulo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        doc.add_paragraph()  # Espacio
        
        # Preparar DataFrame
        df = tabla["data"].copy()
        
        # Limpiar datos vacíos pero mantener filas con al menos un dato
        df = df.dropna(how='all')  # Eliminar filas completamente vacías
        
        # Intentar limpiar celdas que solo tienen espacios o caracteres extraños
        for col in df.columns:
            df[col] = df[col].apply(lambda x: x if pd.notna(x) and str(x).strip() != '' else pd.NA)
        
        # Eliminar filas que quedaron completamente vacías después de la limpieza
        df = df.dropna(how='all')
        
        if len(df) == 0:
            print(f"      ⚠️  Tabla '{titulo_tabla}' sin datos después de limpieza")
            doc.add_paragraph(f"(Tabla encontrada pero sin datos válidos: {titulo_tabla})")
            doc.add_paragraph()
            continue
        
        print(f"      ✅ Tabla con {len(df)} filas válidas")
        
        # Traducir encabezados si corresponde
        if idioma.startswith("es"):
            df.columns = [TRADUCCION_HEADERS.get(col, {}).get("es", col) for col in df.columns]
        else:
            df.columns = [TRADUCCION_HEADERS.get(col, {}).get("en", col) for col in df.columns]
        
        # Crear tabla en Word
        ncols = len(df.columns)
        nrows = len(df)
        tabla_word = doc.add_table(rows=nrows+1, cols=ncols)
        tabla_word.autofit = False
        
        # Configurar encabezados
        for j, col in enumerate(df.columns):
            cell = tabla_word.cell(0, j)
            cell.text = str(col)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0] if p.runs else p.add_run()
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.name = "Arial"
            
            # Añadir sombreado gris a los encabezados
            shading_elm = parse_xml(
                r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="D9D9D9"/>'
            )
            cell._tc.get_or_add_tcPr().append(shading_elm)
        
        # Llenar datos
        for i, (_, row) in enumerate(df.iterrows()):
            for j, val in enumerate(row):
                cell = tabla_word.cell(i+1, j)
                # Formatear el valor apropiadamente
                if pd.isna(val):
                    cell.text = ""
                else:
                    # Si es numérico, formatear con pocas decimales
                    if isinstance(val, (int, float)):
                        if val == int(val):
                            cell.text = str(int(val))
                        else:
                            cell.text = f"{val:.3f}".rstrip('0').rstrip('.')
                    else:
                        cell.text = str(val).strip()
                
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.runs[0] if p.runs else p.add_run()
                run.font.size = Pt(10)
                run.font.name = "Arial"
        
        # Aplicar bordes
        poner_bordes_tabla(tabla_word)
        
        doc.add_paragraph()  # Espacio entre tablas
    
    return doc

def main():
    """Función principal"""
    print("=== EXTRACTOR DE TABLAS DE STAAD PARA CONSOLA ===")
    
    # Seleccionar archivo Excel
    excel_path = select_excel_file()
    if not excel_path:
        print("❌ No se seleccionó archivo.")
        return
    
    # Seleccionar hoja
    sheet_name = select_sheet_name(excel_path)
    if not sheet_name:
        print("❌ No se pudo seleccionar una hoja.")
        return
    
    # Seleccionar idioma
    idioma_input = input("\nSelecciona idioma (es/en) [ENTER=es]: ").strip().lower()
    idioma = "es" if idioma_input in ["", "es", "español"] else "en"
    
    print(f"\n📄 Procesando: {os.path.basename(excel_path)} | Hoja: {sheet_name}")
    
    # Extraer tablas
    tablas_extraidas = extract_target_tables(excel_path, sheet_name=sheet_name)
    
    if not tablas_extraidas:
        print("❌ No se encontraron tablas que coincidan con los patrones definidos.")
        return
    
    print(f"✅ Se encontraron {len(tablas_extraidas)} tablas\n")
    
    # Mostrar tablas en consola
    for idx, tabla in enumerate(tablas_extraidas, 1):
        titulo = TITULOS_TABLAS[tabla["pattern_index"]]["es"] if idioma.startswith("es") else TITULOS_TABLAS[tabla["pattern_index"]]["en"]
        print(f"\n{'='*60}\n{idx}. {titulo} ({len(tabla['data'])} filas)")
        print(f"   Filas {tabla['start_row']} a {tabla['end_row']} en Excel")
        # Traducir encabezados si corresponde
        df = tabla["data"].copy()
        if idioma.startswith("es"):
            df.columns = [TRADUCCION_HEADERS.get(col, {}).get("es", col) for col in df.columns]
        else:
            df.columns = [TRADUCCION_HEADERS.get(col, {}).get("en", col) for col in df.columns]
        print(df.to_string(index=False))
        print('\n')
        
if __name__ == "__main__":
    main()